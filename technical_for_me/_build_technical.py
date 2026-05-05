#!/usr/bin/env python3
"""
_build_technical.py -- Generate one .docx per script for the technical_for_me/
directory.

Each .docx is a deep, internal-only "dissection" of one script:
    1.  Purpose
    2.  Inputs and outputs (data structures)
    3.  Parameter reference (every flag, env-var, and decision rule)
    4.  Algorithm walkthrough (line-by-line algorithmic intent)
    5.  Flow diagram
    6.  Edge cases and known caveats
    7.  Implementation notes (Bash 3.2 tricks, conda env handling, etc.)
    8.  Testing checklist

The contents are kept abstract -- no real run-time paths.

Run:
    cd mbXPro/technical_for_me
    python3 _build_technical.py
"""

from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
HERE = Path(__file__).parent
sys.path.insert(0, str(HERE))
from _flowchart import render_flowchart, render_legend, NAVY, BLUE, GREY  # noqa


DIAGRAMS_DIR = HERE / "_diagrams"
DIAGRAMS_DIR.mkdir(parents=True, exist_ok=True)
LEGEND_PNG = DIAGRAMS_DIR / "_legend.png"
LOGO = HERE.parent / "documentation" / "images" / "mbX_Pro_icon.png"


# =============================================================================
#  Reusable docx styling helpers
# =============================================================================

def set_default_font(doc, name="Calibri", size=11):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return p


def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)
    return p


def add_para(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F4F7")
    pPr.append(shd)
    return p


def add_kv(doc, key, val):
    p = doc.add_paragraph()
    r = p.add_run(f"{key}: ")
    r.bold = True
    p.add_run(str(val))
    return p


def add_table(doc, header, rows, widths_cm=None, fontsize=10):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = h
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(fontsize)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            c = t.cell(i, j)
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(fontsize)
    if widths_cm:
        for row in t.rows:
            for j, w in enumerate(widths_cm):
                if j < len(row.cells):
                    row.cells[j].width = Cm(w)
    return t


def add_logo_centered(doc, width_in=2.0):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(LOGO), width=Inches(width_in))


def add_image_centered(doc, png_path, width_in=6.2, caption=None):
    if not Path(png_path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(png_path), width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x5A, 0x64, 0x70)


def add_page_break(doc):
    doc.add_page_break()


# =============================================================================
#  Chapter scaffolding -- common preamble + cover
# =============================================================================

def make_doc_with_cover(title, subtitle, script_name, step_id):
    """Initialize a Document with the standard cover page."""
    doc = Document()
    set_default_font(doc, "Calibri", 11)

    add_logo_centered(doc, width_in=1.6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("mbX Pro -- internal technical reference")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x5A, 0x64, 0x70)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Step {step_id}")
    r.font.size = Pt(13)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(script_name)
    r.font.name = "Courier New"
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x5A, 0x64, 0x70)

    return doc


def add_legend(doc):
    if not LEGEND_PNG.exists():
        render_legend(LEGEND_PNG)
    add_image_centered(doc, LEGEND_PNG, width_in=5.5,
                       caption="Flowchart symbol legend (used in every chapter).")


# =============================================================================
#  Standard section helpers
# =============================================================================

def section_purpose(doc, text):
    add_h1(doc, "1. Purpose and scope")
    add_para(doc, text)


def section_inputs_outputs(doc, inputs_table, outputs_table, upstream_consumers,
                           input_caption=None, output_caption=None):
    add_h1(doc, "2. Inputs and outputs")
    add_h2(doc, "2.1 Inputs (data structures)")
    if input_caption:
        add_para(doc, input_caption)
    add_table(doc,
              header=["Name", "Source", "Format", "Required"],
              rows=inputs_table,
              widths_cm=[3.5, 4.0, 5.5, 2.0],
              fontsize=9)

    add_h2(doc, "2.2 Outputs (data structures)")
    if output_caption:
        add_para(doc, output_caption)
    add_table(doc,
              header=["Path (relative to output dir)", "What it contains", "Read by"],
              rows=outputs_table,
              widths_cm=[5.5, 6.0, 3.5],
              fontsize=9)

    add_h2(doc, "2.3 Downstream consumers")
    add_para(doc, "These steps depend on this script's outputs:")
    for c in upstream_consumers:
        add_bullet(doc, c)


def section_parameters(doc, params_table, intro=None):
    add_h1(doc, "3. Complete parameter reference")
    if intro:
        add_para(doc, intro)
    add_table(doc,
              header=["Parameter", "Default", "Decision rule / where it comes from"],
              rows=params_table,
              widths_cm=[4.5, 2.5, 8.0],
              fontsize=9)


def section_algorithm(doc, walkthrough):
    """walkthrough is a list of (subsection_title, paragraphs_or_code_or_bullets) tuples.
    Each item in the second slot is a tuple (kind, content):
      kind='p' -> add_para(content)
      kind='code' -> add_code(content)
      kind='bullets' -> list of bullet strings
      kind='kv' -> list of (key, value) tuples for kv lines"""
    add_h1(doc, "4. Algorithm walkthrough")
    for title, items in walkthrough:
        add_h2(doc, title)
        for kind, content in items:
            if kind == "p":
                add_para(doc, content)
            elif kind == "code":
                add_code(doc, content)
            elif kind == "bullets":
                for b in content:
                    add_bullet(doc, b)
            elif kind == "kv":
                for k, v in content:
                    add_kv(doc, k, v)
            elif kind == "h3":
                add_h3(doc, content)


def section_flow(doc, png_path, caption):
    add_h1(doc, "5. Flow diagram")
    add_legend(doc)
    add_image_centered(doc, png_path, width_in=6.5, caption=caption)


def section_edge_cases(doc, items):
    add_h1(doc, "6. Edge cases and caveats")
    for title, body in items:
        add_h3(doc, title)
        add_para(doc, body)


def section_impl_notes(doc, items):
    add_h1(doc, "7. Implementation notes")
    for title, body in items:
        add_h3(doc, title)
        if isinstance(body, list):
            for line in body:
                add_para(doc, line)
        else:
            add_para(doc, body)


def section_testing(doc, items):
    add_h1(doc, "8. Testing checklist")
    add_para(doc, "Things to verify before considering this script production-ready:")
    for it in items:
        add_bullet(doc, it)


# =============================================================================
#  Per-script chapter builders.
#  Each builder produces ONE .docx and the matching flowchart PNG.
# =============================================================================

def build_01_primer_identifier():
    NAME = "1_mbx_primer_identifier.docx"
    doc = make_doc_with_cover(
        "Auto-detect 16S rRNA primers from raw FASTQ files",
        "Pipeline step 0 -- the very first thing that runs.",
        "mbx_primer_identifier.sh",
        step_id="0",
    )

    section_purpose(doc, (
        "Determine which forward and reverse primers were used during library "
        "preparation by sampling a representative number of reads and matching "
        "them against a built-in database of 30 published 16S rRNA primers. "
        "The result drives the trim-left parameters of step 4 (DADA2). "
        "If primers cannot be identified above the user-tunable rate, the "
        "pipeline falls back to default trim-left=20 and the user can pass "
        "--forward-primer / --reverse-primer manually."
    ))

    section_inputs_outputs(doc,
        inputs_table=[
            ["<fastq_dir>", "User CLI argument", "Directory containing *.fastq.gz files (paired or single-end)", "yes"],
            ["--samples N", "User flag", "Integer (default 10000)", "no"],
            ["--mismatches N", "User flag", "Integer 0-5 (default 2)", "no"],
            ["--offset N", "User flag", "Integer 0-30 (default 15)", "no"],
            ["--min-rate F", "User flag", "Float 0-1 (default 0.10)", "no"],
            ["--no-search-parent / --no-search-child", "User flag", "Boolean toggle", "no"],
            ["MBX_OUT_DIR", "Env var (set by orchestrator)", "Absolute path", "no"],
        ],
        outputs_table=[
            ["0_primer_handling/mbx_primer_info.txt",
             "Key=value file with detected primers + diagnostic data",
             "step 5 (classifier_arranger), step 3 (dada2_parameter_finder)"],
        ],
        upstream_consumers=[
            "Step 3 (mbx_dada2_parameter_finder.sh) reads FORWARD_PRIMER_SEQUENCE / REVERSE_PRIMER_SEQUENCE to set trim-left.",
            "Step 5 (mbx_classifier_arranger.sh) reads the same fields to invoke `qiime feature-classifier extract-reads`.",
            "Step 18 (mbx_final_report.sh) reads the entire info file for the report's `Primer detection` section.",
        ],
    )

    section_parameters(doc,
        intro=(
            "The script is parameter-light by design -- the user typically "
            "passes only the FASTQ directory. Every flag has a sensible "
            "default tuned to standard Illumina 16S libraries."
        ),
        params_table=[
            ["--samples N",       "10000",
             "Number of reads to sample PER FILE. Reads are taken from the start of the file (NOT randomly) -- this is intentional because Illumina FASTQ records do not have temporal order to worry about, and sequential sampling is dramatically faster than random seek over gzip."],
            ["--mismatches N",    "2",
             "Maximum allowed IUPAC-aware mismatches between read prefix and candidate primer. 2 is generous for V3-V4 (19-20 bp primers) and matches Cutadapt's typical 0.1 error-rate setting. Increase to 3 if detection fails."],
            ["--offset N",        "15",
             "Maximum left-shift of the primer when sliding it along the read. Some sequencing facilities leave a few bases of adapter or random padding before the primer. 15 covers every protocol we have encountered."],
            ["--min-rate F",      "0.10",
             "Minimum fraction of sampled reads that must match the same primer for it to be reported. Below this fraction, the primer is reported as 'None'. 0.10 = 10% is intentionally conservative -- a true primer pair typically matches >70%."],
            ["MBX_OUT_DIR (env)", "(unset)",
             "When set by the orchestrator, the script reuses that exact directory as OUT_ROOT and derives TIMESTAMP from its name. When unset, the script generates its own timestamp and creates a sibling output directory."],
        ],
    )

    section_algorithm(doc, [
        ("4.1 Preflight", [
            ("p", "After argument parsing the script verifies that python3 is on PATH and the FASTQ directory exists. python3 is a hard dependency because all heavy lifting is delegated to an embedded Python heredoc."),
            ("p", "MBX_OUT_DIR is checked. If set, OUT_ROOT is reused directly; otherwise OUT_ROOT becomes <parent_of_FASTQ>/mbX_pro_outputs_<TIMESTAMP>. This contract is what guarantees that the orchestrator's single-shared-directory invariant holds."),
        ]),
        ("4.2 File discovery (Bash 3.2 safe)", [
            ("p", "Discovery uses three concentric search rings with deduplication via a temp file (no associative arrays). Each candidate directory is added to dirs.txt only if its absolute path is not already present (grep -qxF)."),
            ("bullets", [
                "Ring 1: <fastq_dir> itself (always).",
                "Ring 2: parent directory of <fastq_dir> (unless --no-search-parent).",
                "Ring 3: every direct child directory of <fastq_dir> (unless --no-search-child).",
            ]),
            ("p", "Within each ring, `find -maxdepth 1 -iname '*.fastq.gz'` lists files. Each file's absolute path is compared against seen.txt to deduplicate. The result is files.txt."),
        ]),
        ("4.3 R1 / R2 split", [
            ("p", "Each filename is matched case-insensitively against 'r1' and 'r2' (regex). Files matching neither tag are skipped with a warning. The presence of any R2 file flips READ_TYPE to 'paired'; otherwise 'single'."),
        ]),
        ("4.4 Python primer engine -- IUPAC table", [
            ("p", "The Python heredoc starts by defining the IUPAC expansion dictionary. Each ambiguous base maps to a frozenset of unambiguous bases:"),
            ("code",
                "M -> {A,C}     R -> {A,G}     W -> {A,T}\n"
                "S -> {C,G}     Y -> {C,T}     K -> {G,T}\n"
                "B -> {C,G,T}   D -> {A,G,T}   H -> {A,C,T}\n"
                "V -> {A,C,G}   N -> {A,C,G,T}\n"
                "U -> {T}  (RNA -> DNA)\n"
                "Lowercase keys mirror uppercase keys for tolerant matching."),
            ("p", "An RC_TABLE for str.translate then implements reverse complementation in O(n) without dictionary lookups."),
        ]),
        ("4.5 Built-in primer database", [
            ("p", "PRIMER_DB is a list of dicts (name, seq, region, dir). 30 published 16S primer entries cover V1-V2, V1-V3, V1-V9, V3, V3-V4 (incl. PRK universal), V3-V5, V4 (original + EMP-Parada/Apprill), V4-V5, V5-V7, V6-V7, V6-V8."),
            ("p", "After defining the DB, every sequence is scrubbed: the Cyrillic look-alike `\\u041d` (Cyrillic Capital En) -> ASCII 'N' (a copy-paste hazard from primer Excel sheets), and uppercased. MIN_PRIMER_LEN and MAX_PRIMER_LEN are computed at module load -- used later as the consensus diagnostic window width."),
        ]),
        ("4.6 The match function", [
            ("p", "iupac_mismatches(read_prefix, primer) walks both sequences in lockstep. At each position, the script checks whether the read base belongs to IUPAC[primer_base]; if not, increments mm. Returns the total mismatch count."),
            ("p", "slide_match(read, primer, max_mm, max_offset) tries every left-shift from 0 to max_offset. For each shift, it slices read[offset:offset+len(primer)] and calls iupac_mismatches. The minimum mismatch (and the offset that achieved it) is tracked. Returns (matched, best_offset). matched is True iff best_mm <= max_mm. A 0-mismatch hit short-circuits the loop."),
        ]),
        ("4.7 Scoring all primers", [
            ("p", "score_primers(reads, primers, max_mm, max_offset, as_rc) returns a dict keyed by primer name. Each entry records: matched count, total reads scanned, rate (matched/total), modal_offset (the most common offset that produced a hit), and as_rc flag."),
            ("p", "Two passes are run on the R1 reads -- one with as_rc=False (testing primer 5'->3' as written) and one with as_rc=True (testing reverse-complement). For paired-end runs the same two passes are run on R2 reads."),
        ]),
        ("4.8 Three-orientation hypothesis testing (key innovation)", [
            ("p", "Different sequencing facilities and chemistries produce three subtly different read orientations. The script tests all three and picks the highest combined match rate:"),
            ("kv", [
                ("Orientation A (RC)",      "R1 starts with forward primer 5'->3'; R2 starts with reverse-complement of reverse primer."),
                ("Orientation B (DIRECT)",  "R1 starts with forward primer 5'->3'; R2 starts with reverse primer 5'->3' (no RC)."),
                ("Orientation C (SWAP)",    "R1 starts with reverse primer; R2 starts with forward primer (some facilities label R1/R2 swapped)."),
            ]),
            ("p", "Each orientation's score is the sum of best-forward-primer rate + best-reverse-primer rate. Whichever orientation has the highest sum wins. orientation_note records which orientation was selected for the report."),
        ]),
        ("4.9 Consensus diagnostic", [
            ("p", "consensus_bases(reads, length) computes the per-position modal base across all sampled reads. The window length is dynamically sized to MAX_PRIMER_LEN + max_offset (e.g. 22 + 15 = 37 nt). This fixes a former bug where the window was hard-coded to 40 nt and could mis-diagnose short primers."),
            ("p", "The consensus is printed both to stdout (prefixed [DIAG]) and embedded in the output file, so a user reading mbx_primer_info.txt can see exactly what the read prefix looks like."),
        ]),
        ("4.10 Output formatting", [
            ("p", "fmt_primer(label, best, direction_str) returns a list of key=value strings. When best is None, the function emits a comment block explaining how to relax thresholds, plus the same key=None lines so downstream parsers do not fail."),
            ("p", "The full output file contains: a 10-line metadata header (date, FASTQ dir, sampling counts, mismatches, offset, min-rate, orientation, consensus); the READ_TYPE line; FORWARD_PRIMER_* block (7 lines); REVERSE_PRIMER_* block (7 lines for paired; 1 line for single)."),
        ]),
    ])

    # Flow diagram
    flow_png = DIAGRAMS_DIR / "01_primer_identifier.png"
    nodes = [
        dict(id="start", x=5, y=11.0, w=2.6, h=0.7, label="START\nmbx_primer_identifier.sh <fastq_dir>", kind="term"),
        dict(id="parse", x=5, y=10.0, w=4.6, h=0.7, label="Parse args + verify python3, FASTQ dir"),
        dict(id="outdir", x=5, y=9.0, w=5.0, h=0.8, label="MBX_OUT_DIR set?\nyes -> reuse it    no -> generate timestamped dir", kind="decision"),
        dict(id="discover", x=5, y=7.9, w=5.5, h=0.7, label="Discover *.fastq.gz files\n(self + parent + 1-level child)"),
        dict(id="split", x=5, y=6.9, w=4.6, h=0.7, label="Split by R1/R2 tag (case-insensitive)"),
        dict(id="sample", x=5, y=5.9, w=5.5, h=0.7, label="Sample N reads per file (round-robin)"),
        dict(id="score", x=5, y=4.9, w=6.4, h=0.8, label="Score 30 primers x 2 orientations\n(IUPAC sliding window 0..max_offset)"),
        dict(id="orient", x=5, y=3.6, w=6.4, h=1.0, label="Pick best orientation (RC / DIRECT / SWAP)\nbased on combined fwd+rev rate", kind="decision"),
        dict(id="thresh", x=5, y=2.4, w=6.0, h=0.8, label="rate >= min_rate?\nyes -> emit primer    no -> emit 'None'", kind="decision"),
        dict(id="write", x=5, y=1.3, w=5.5, h=0.7, label="Write mbx_primer_info.txt", kind="io"),
        dict(id="end", x=5, y=0.4, w=2.0, h=0.6, label="END", kind="term"),
    ]
    edges = [
        dict(**{"from": "start", "to": "parse"}),
        dict(**{"from": "parse", "to": "outdir"}),
        dict(**{"from": "outdir", "to": "discover"}),
        dict(**{"from": "discover", "to": "split"}),
        dict(**{"from": "split", "to": "sample"}),
        dict(**{"from": "sample", "to": "score"}),
        dict(**{"from": "score", "to": "orient"}),
        dict(**{"from": "orient", "to": "thresh"}),
        dict(**{"from": "thresh", "to": "write"}),
        dict(**{"from": "write", "to": "end"}),
    ]
    render_flowchart(nodes, edges, flow_png,
                     title="mbx_primer_identifier.sh -- control flow")
    section_flow(doc, flow_png, "Control flow of mbx_primer_identifier.sh.")

    section_edge_cases(doc, [
        ("Pre-trimmed primers",
         "Many sequencing facilities trim primers before delivery. In that case all primers will fall below --min-rate and the script reports 'None'. The user must then pass --forward-primer/--reverse-primer manually OR live with the trim-left=20 default."),
        ("Very short reads",
         "Reads shorter than the longest primer + max_offset are skipped silently in slide_match (the window slice is empty)."),
        ("Cyrillic look-alikes in the primer DB",
         "After the database literal definition, every sequence is post-processed to replace Cyrillic Capital En (\\u041d) with ASCII 'N'. This is a defensive measure against copy-paste from primer Excel sheets."),
        ("Single-end data",
         "READ_TYPE='single' branches at multiple points: only one consensus is computed, only the forward primer is reported, and the REVERSE_PRIMER_NAME line shows 'N/A (single-end)' so downstream parsers can detect this case."),
        ("Match-rate ties",
         "If two primers tie on rate, max() returns whichever appears first in PRIMER_DB. The DB is ordered with EMP-Parada/Apprill listed BEFORE the original 515F/806R variants, so a tie favours the more current EMP version."),
    ])

    section_impl_notes(doc, [
        ("PID-based temp directory cleanup",
         "trap 'rm -rf \"$WORK_DIR\"' EXIT ensures the temp work-dir is removed on every exit path including ctrl-C. WORK_DIR is generated via mktemp -d, which is safe (only the .R suffix mktemp variant has the macOS XXXXXX bug)."),
        ("Environment-var contract with the Python heredoc",
         "Twelve MBX_* environment variables are exported before the heredoc. This avoids quoting and shell-injection problems entirely -- the Python script reads os.environ directly. The heredoc is enclosed in single-quoted PYEOF markers so $-expansion does not trigger inside the Python source."),
        ("Why sliding-window instead of exact prefix",
         "Exact prefix matching at offset 0 fails as soon as ANY adapter or padding base is present. Sliding 0..max_offset covers every protocol with a single uniform algorithm. The performance cost is negligible because each read's window is only 22 + 15 = 37 nt and only 30 primer candidates are tested."),
    ])

    section_testing(doc, [
        "Run with a known-primer FASTQ set (515F/806R EMP) and confirm match rate > 70 %.",
        "Run with --no-search-parent and verify only the FASTQ directory is searched.",
        "Run with --mismatches 0 -- detection should still succeed with primer-pristine data.",
        "Run with primers already trimmed -- confirm the script reports 'None' AND prints the 'pre-trimmed' tip.",
        "Run with the orchestrator setting MBX_OUT_DIR -- confirm output lands in MBX_OUT_DIR/0_primer_handling/, NOT a fresh timestamped dir.",
        "Confirm that `qiime feature-classifier extract-reads` accepts the resulting primer sequences without modification.",
    ])

    doc.save(HERE / NAME)
    print(f"[OK] {NAME}  ({(HERE / NAME).stat().st_size/1024:.1f} KB)")


def build_02_create_manifest():
    NAME = "2_create_manifest.docx"
    doc = make_doc_with_cover(
        "Build a QIIME2-format paired/single-end manifest TSV",
        "Pipeline step 1 -- the bridge between filesystem and QIIME2.",
        "create_manifest.sh",
        step_id="1",
    )

    section_purpose(doc, (
        "Discover paired R1/R2 (or single-end) FASTQ files inside a "
        "user-supplied directory, extract per-sample IDs, validate that "
        "every paired-end sample has BOTH R1 and R2, and write a "
        "QIIME2-compatible manifest file with absolute filepaths."
    ))

    section_inputs_outputs(doc,
        inputs_table=[
            ["<fastq_dir>", "User CLI argument", "Directory containing *.fastq.gz", "yes"],
            ["--no-search-parent", "User flag", "Boolean toggle", "no"],
            ["--no-search-child", "User flag", "Boolean toggle", "no"],
            ["MBX_OUT_DIR (env)", "Set by orchestrator", "Absolute path", "no"],
        ],
        outputs_table=[
            ["1_manifest_file/manifest.txt",
             "Tab-separated QIIME2 manifest. Header is 'sample-id\\tforward-absolute-filepath\\treverse-absolute-filepath' (paired) or 'sample-id\\tabsolute-filepath' (single-end)",
             "step 2 (artifact_creator.sh)"],
        ],
        upstream_consumers=[
            "Step 2 (artifact_creator.sh) reads the manifest header to detect paired vs single-end and runs `qiime tools import` with the matching --input-format.",
        ],
    )

    section_parameters(doc,
        intro=(
            "Apart from search-ring toggles there are no tunable parameters. "
            "Sample ID extraction is rule-based and deterministic."
        ),
        params_table=[
            ["--no-search-parent", "off",
             "Disables Ring 2 (parent of <fastq_dir>). Useful when the parent contains FASTQ files from an unrelated run."],
            ["--no-search-child",  "off",
             "Disables Ring 3 (children of <fastq_dir>). Useful when the FASTQ directory contains output sub-folders that incidentally hold .fastq.gz files."],
            ["MBX_OUT_DIR (env)", "(unset)",
             "Same contract as step 0: when set by the orchestrator, the manifest lands in MBX_OUT_DIR/1_manifest_file/."],
        ],
    )

    section_algorithm(doc, [
        ("4.1 Bash 3.2-safe associative array replacement", [
            ("p", "macOS Bash 3.2 has no `declare -A`. We need a mapping sample_id -> R1 path and another sample_id -> R2 path. The script implements this with a temp-directory file-based key/value store:"),
            ("code",
                "WORK_DIR=$(mktemp -d)\n"
                "mkdir -p \"$WORK_DIR/R1\" \"$WORK_DIR/R2\"\n"
                "# Each \"slot\" is a file named after the sample ID:\n"
                "echo \"$path\" > \"$WORK_DIR/R1/$sample_id\"\n"
                "# Reading is just `cat`:\n"
                "r1=$(cat \"$WORK_DIR/R1/$sample_id\")"),
            ("p", "Duplicate detection is automatic: if the slot file already exists when the script tries to write, that means two FASTQ files mapped to the same (sample_id, R1|R2) tuple -- a fatal error."),
        ]),
        ("4.2 File discovery", [
            ("p", "Identical algorithm to step 0: three concentric rings (self, parent, immediate children). Files are deduplicated via the seen.txt approach so that overlapping ring memberships do not double-process a file."),
        ]),
        ("4.3 R1/R2 detection", [
            ("p", "_detect_read() runs `grep -qiE 'r1'` then `grep -qiE 'r2'` on the basename. The case-insensitive flag means both Sample_R1.fastq.gz and sample-r1-001.fastq.gz match."),
            ("p", "Files matching neither pattern are added to the UNDETECTED counter and skipped. If PROCESSED reaches 0, the script aborts with an actionable error."),
        ]),
        ("4.4 Sample ID extraction (two-strategy fallback)", [
            ("h3", "Strategy 1: anchored sample regex"),
            ("p", "Match `s(ample)?[-_]?[0-9]+` case-insensitively. This catches every Illumina-naming convention: S1, S1234, Sample-1, sample_42, SAMPLE-99."),
            ("h3", "Strategy 2: filename-cleanup"),
            ("p", "If Strategy 1 returns empty, fall back to a Perl pipeline that strips: .fastq.gz, R1/R2 tags, leading and trailing .-_ characters."),
            ("code",
                "id=$(printf '%s' \"$nm\" | perl -pe '\n"
                "  s/\\.fastq\\.gz$//i;\n"
                "  s/r[12]//gi;\n"
                "  s/[._-]+$//;\n"
                "  s/^[._-]+//;\n"
                "  chomp;\n"
                "')"),
            ("p", "The result is uppercased via `tr '[:lower:]' '[:upper:]'` (avoids the Bash 4-only ${var^^})."),
        ]),
        ("4.5 Validation -- paired-end", [
            ("p", "After processing, the script collects every unique sample ID from R1/ and R2/ via `sort -u`. For each ID, it verifies BOTH R1/$sid and R2/$sid exist; missing IDs accumulate in MISSING[]."),
            ("p", "If MISSING is non-empty, the script aborts with a per-sample list. This catches the most common user error: 'I deleted one of the R2 files by accident'."),
        ]),
        ("4.6 Manifest emission", [
            ("p", "The manifest is written sample-by-sample, alphabetically sorted. The header line is the QIIME2-required format (paired or single). The script ends with a printout of the suggested `qiime tools import` command -- helpful for users running the steps individually."),
        ]),
    ])

    flow_png = DIAGRAMS_DIR / "02_create_manifest.png"
    nodes = [
        dict(id="start", x=5, y=11.0, w=2.4, h=0.7, label="START\ncreate_manifest.sh <fastq_dir>", kind="term"),
        dict(id="parse", x=5, y=9.9, w=4.4, h=0.7, label="Parse args + verify dir"),
        dict(id="outdir", x=5, y=8.8, w=5.0, h=0.8, label="MBX_OUT_DIR set?", kind="decision"),
        dict(id="discover", x=5, y=7.6, w=5.0, h=0.7, label="Discover *.fastq.gz (3 rings)"),
        dict(id="split", x=5, y=6.5, w=5.0, h=0.8, label="Detect R1/R2 + extract sample-id\n(two-strategy fallback)"),
        dict(id="kvstore", x=5, y=5.2, w=5.6, h=0.9, label="Write to temp KV store:\nWORK_DIR/R1/<sid>  +  WORK_DIR/R2/<sid>", kind="io"),
        dict(id="dups", x=5, y=4.0, w=4.6, h=0.8, label="Slot exists already?\n(duplicate detection)", kind="decision"),
        dict(id="paired", x=5, y=2.7, w=4.6, h=0.8, label="Any R2 file present?\nyes -> paired   no -> single", kind="decision"),
        dict(id="missing", x=5, y=1.5, w=5.6, h=0.8, label="(paired only) verify every sid\nhas both R1 and R2", kind="decision"),
        dict(id="write", x=5, y=0.4, w=4.4, h=0.5, label="Write manifest.txt", kind="io"),
    ]
    edges = [
        dict(**{"from": "start", "to": "parse"}),
        dict(**{"from": "parse", "to": "outdir"}),
        dict(**{"from": "outdir", "to": "discover"}),
        dict(**{"from": "discover", "to": "split"}),
        dict(**{"from": "split", "to": "kvstore"}),
        dict(**{"from": "kvstore", "to": "dups"}),
        dict(**{"from": "dups", "to": "paired"}),
        dict(**{"from": "paired", "to": "missing"}),
        dict(**{"from": "missing", "to": "write"}),
    ]
    render_flowchart(nodes, edges, flow_png, title="create_manifest.sh -- control flow")
    section_flow(doc, flow_png, "Control flow of create_manifest.sh.")

    section_edge_cases(doc, [
        ("Mixed-case R1/R2 in same run",
         "Some facilities use 'R1', others 'r1'. The case-insensitive grep handles both."),
        ("Numeric sample IDs",
         "The regex `s[-_]?[0-9]+` matches s1 / s_1 / s-1. Plain numeric IDs like '42.fastq.gz' fall to Strategy 2 cleanup, which yields '42'."),
        ("Lane-stamped Illumina names",
         "S1_L001_R1_001.fastq.gz -> Strategy 1 matches 'S1' before the lane suffix; the lane and read-number suffixes are silently dropped."),
        ("BOM characters in pasted filenames",
         "Not currently stripped here -- this is handled later in step 4 metadata validation. If a filename starts with U+FEFF the regex still matches."),
        ("Symbolic links",
         "Symlinks are followed transparently because `find -L` is the default for `find` in macOS. The absolute path written to the manifest is the symlink path, not the target -- this is what the user asked for and is what QIIME2 accepts."),
    ])

    section_impl_notes(doc, [
        ("Why two strategies?",
         "Strategy 1 (`s[-_]?[0-9]+`) is precise but misses non-conventional names like 'CTRL_A.fastq.gz'. Strategy 2 (cleanup) catches those but is sometimes too greedy. Running them in order gives precision-first, recall-fallback behaviour."),
        ("Sorting",
         "`sort -u` on the slot filenames provides both deduplication and deterministic ordering. The manifest is therefore byte-identical when the same FASTQ set is re-processed."),
        ("Why no associative array",
         "`declare -A` is Bash 4+. macOS ships Bash 3.2.57 by default. The temp-directory KV store is the standard workaround documented in CLAUDE.md and used by every script in this pipeline that needs an associative mapping."),
    ])

    section_testing(doc, [
        "Run with paired-end Illumina data and confirm the manifest header has 3 columns.",
        "Run with single-end data (only R1 files) and confirm the header has 2 columns.",
        "Manually delete one R2 file before running and confirm the script aborts with the missing-file error.",
        "Place two R1 files for the same sample-id in different sub-folders and confirm the duplicate-detection error fires.",
        "Verify that `qiime tools import` accepts the resulting manifest without modification.",
    ])

    doc.save(HERE / NAME)
    print(f"[OK] {NAME}  ({(HERE / NAME).stat().st_size/1024:.1f} KB)")


def build_03_artifact_creator():
    NAME = "3_artifact_creator.docx"
    doc = make_doc_with_cover(
        "Import the manifest into a QIIME2 .qza artifact + summary .qzv",
        "Pipeline step 2 -- the first pure-QIIME2 wrapper.",
        "artifact_creator.sh",
        step_id="2",
    )

    section_purpose(doc, (
        "Run `qiime tools import` on the manifest produced by step 1, then "
        "`qiime demux summarize` to produce a quality-summary visualization. "
        "The script automatically detects paired vs single-end from the "
        "manifest header so the user does not have to set a flag."
    ))

    section_inputs_outputs(doc,
        inputs_table=[
            ["<manifest.txt>", "From step 1", "Tab-separated, sample-id + filepaths", "yes"],
            ["--dry-run", "User flag", "Boolean toggle", "no"],
        ],
        outputs_table=[
            ["2_first_artifact_file/Paired_End_artifact.qza",
             "QIIME2 SampleData[PairedEndSequencesWithQuality] artifact (or Single_End_artifact.qza for single-end)",
             "step 3 (mbx_dada2_parameter_finder.sh)"],
            ["2_first_artifact_file/Paired_End_artifact.qzv",
             "Visualization with per-position read-quality summary",
             "step 3 (mbx_dada2_parameter_finder.sh)"],
        ],
        upstream_consumers=[
            "Step 3 (mbx_dada2_parameter_finder.sh) reads both .qza and .qzv to extract per-position quality statistics for trunc-len selection.",
            "Step 4 (mbx_dada2_run.sh) operates on the .qza directly via the path stored in dada2_parameters.txt.",
        ],
    )

    section_parameters(doc,
        intro="The script is parameter-free. All decisions are derived from the manifest header.",
        params_table=[
            ["--dry-run", "off",
             "Print every command that WOULD be executed but skip the actual run. Useful when the user wants to inspect commands before committing to a slow run."],
        ],
    )

    section_algorithm(doc, [
        ("4.1 Pre-flight checks", [
            ("bullets", [
                "Verify <manifest.txt> exists.",
                "Verify `qiime` is on PATH (catches the most common user mistake -- forgetting to activate the conda env).",
                "Print `qiime info` version line for provenance.",
                "Validate that the manifest has at least one data row (header-only files indicate a failed step 1).",
                "Walk the manifest and verify every filepath listed exists -- catches 'data on USB drive that is not mounted' before QIIME2 produces a less-helpful error.",
            ]),
        ]),
        ("4.2 Read-type detection", [
            ("p", "The first line of the manifest is read. If it contains `forward-absolute-filepath` -> paired-end; if only `absolute-filepath` -> single-end. Anything else triggers a fatal error with the expected header format printed."),
        ]),
        ("4.3 Output-directory derivation", [
            ("p", "The manifest path is expected to look like `.../mbX_pro_outputs_*/1_manifest_file/manifest.txt`. The script derives the parent of `1_manifest_file/` as `MBX_OUT_DIR` (note: this is a LOCAL variable inside this script -- not the env-var contract used by step 0/1). Then it creates `MBX_OUT_DIR/2_first_artifact_file/`."),
            ("p", "If the manifest's parent folder is NOT named `1_manifest_file/`, a warning fires but the script proceeds. This permits running step 2 stand-alone on a hand-crafted manifest."),
        ]),
        ("4.4 QIIME2 invocation", [
            ("p", "Two QIIME2 commands run in sequence:"),
            ("kv", [
                ("Step 1", "qiime tools import --type SampleData[PairedEndSequencesWithQuality] (or single-end variant) --input-format PairedEndFastqManifestPhred33V2 --input-path <manifest> --output-path <qza>"),
                ("Step 2", "qiime demux summarize --i-data <qza> --o-visualization <qzv>"),
            ]),
            ("p", "If step 1 fails, step 2 is skipped and the user gets the QIIME2 error verbatim plus a 4-bullet 'common causes' list (missing files, corrupted gzip, wrong Phred encoding, disk full)."),
            ("p", "If step 2 fails after step 1 succeeded, the user is told the .qza is OK and given a copy-paste retry command for just the summarize call."),
        ]),
    ])

    flow_png = DIAGRAMS_DIR / "03_artifact_creator.png"
    nodes = [
        dict(id="start", x=5, y=11.0, w=2.4, h=0.7, label="START\nartifact_creator.sh <manifest>", kind="term"),
        dict(id="parse", x=5, y=10.0, w=4.0, h=0.6, label="Parse args + verify manifest"),
        dict(id="qiime", x=5, y=9.0, w=4.0, h=0.6, label="`qiime` on PATH?", kind="decision"),
        dict(id="rt", x=5, y=8.0, w=5.0, h=0.7, label="Read manifest header\nDetect paired vs single-end"),
        dict(id="paths", x=5, y=6.9, w=5.4, h=0.8, label="Validate every filepath in manifest exists\n(catches unmounted drives early)"),
        dict(id="outdir", x=5, y=5.7, w=5.4, h=0.8, label="Derive output dir\nMBX_OUT_DIR/2_first_artifact_file/"),
        dict(id="import", x=5, y=4.5, w=5.4, h=0.8, label="qiime tools import\n(produces .qza)", kind="io"),
        dict(id="failed", x=8.0, y=3.5, w=2.5, h=0.8, label="Print 4-bullet\n'common causes' list", kind="fail"),
        dict(id="summ", x=5, y=3.5, w=4.0, h=0.7, label="qiime demux summarize\n(produces .qzv)", kind="io"),
        dict(id="end", x=5, y=2.4, w=2.0, h=0.6, label="END", kind="term"),
    ]
    edges = [
        dict(**{"from": "start", "to": "parse"}),
        dict(**{"from": "parse", "to": "qiime"}),
        dict(**{"from": "qiime", "to": "rt"}),
        dict(**{"from": "rt", "to": "paths"}),
        dict(**{"from": "paths", "to": "outdir"}),
        dict(**{"from": "outdir", "to": "import"}),
        dict(**{"from": "import", "to": "summ"}, label="ok"),
        dict(**{"from": "import", "to": "failed"}, label="fail"),
        dict(**{"from": "summ", "to": "end"}),
    ]
    render_flowchart(nodes, edges, flow_png, title="artifact_creator.sh -- control flow")
    section_flow(doc, flow_png, "Control flow of artifact_creator.sh.")

    section_edge_cases(doc, [
        ("Manifest with zero data rows",
         "An empty manifest is a sign of a failed step 1. The script aborts before invoking qiime, which would produce a confusing error."),
        ("Phred33 vs Phred64",
         "PairedEndFastqManifestPhred33V2 / SingleEndFastqManifestPhred33V2 is hard-coded. Modern Illumina output is always Phred33. Phred64 (Illumina <1.8, 2011 and earlier) is not supported."),
        (".qzv generation can fail when .qza succeeded",
         "The summarize step needs slightly more memory. If only .qzv fails, the user can re-run just `qiime demux summarize` on the existing .qza."),
    ])

    section_impl_notes(doc, [
        ("Why detect read-type from manifest header (not from filenames)",
         "The manifest is the canonical 'this is what the user actually has' declaration. If the user hand-edited the manifest to remove R2, we trust that and import as single-end."),
        ("File-existence pre-check",
         "QIIME2's own error message for a missing file is buried inside Python tracebacks. We surface it earlier with a per-file warning + summary error."),
    ])

    section_testing(doc, [
        "Run on a paired-end manifest -> verify Paired_End_artifact.qza is produced.",
        "Run on a single-end manifest -> verify Single_End_artifact.qza is produced.",
        "Hand-corrupt one FASTQ file (truncate gzip) -> verify the script's error message includes 'corrupted or not properly gzipped'.",
        "Run with --dry-run -> verify no output files are written.",
    ])

    doc.save(HERE / NAME)
    print(f"[OK] {NAME}  ({(HERE / NAME).stat().st_size/1024:.1f} KB)")


# Continue with the next builders in the next file -- but for now we'll
# expose what we have plus a top-level main() that calls everything.
def main():
    builders = [
        build_01_primer_identifier,
        build_02_create_manifest,
        build_03_artifact_creator,
    ]
    for b in builders:
        b()


if __name__ == "__main__":
    main()
