#!/usr/bin/env python3
"""
_build_docs.py -- Generate the three .docx deliverables for mbX Pro:

   1. mbXPro_documentation.docx   -- Comprehensive technical reference
   2. how_to_run_mbX_Pro.docx     -- End-user installation + run manual
   3. how_to_upload_mbX_Pro.docx  -- GitHub publishing guide for the maintainer

This script is part of the build pipeline.  Run it from inside the
documentation/ directory:
   cd mbXPro/documentation
   python3 _build_docs.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = Path(__file__).parent
LOGO = HERE / "images" / "mbX_Pro_icon.png"


# =============================================================================
#  Reusable styling helpers
# =============================================================================

def set_default_font(doc, name="Calibri", size=11):
    """Set Normal style font + size."""
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return p


def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)
    return p


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def add_kv_para(doc, key, value):
    p = doc.add_paragraph()
    r = p.add_run(f"{key}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    p.add_run(str(value))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def add_code_block(doc, text):
    """Add a monospace, lightly-shaded code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9.5)
    # Apply a light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F4F7")
    pPr.append(shd)
    return p


def add_table_2col(doc, rows):
    """Two-column table: rows is a list of (header, value) tuples."""
    t = doc.add_table(rows=len(rows), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        t.cell(i, 0).text = k
        t.cell(i, 1).text = v
        for run in t.cell(i, 0).paragraphs[0].runs:
            run.bold = True
    # Set column widths
    for row in t.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(10.5)
    return t


def add_logo_centered(doc, width_in=2.4):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(LOGO), width=Inches(width_in))


def add_page_break(doc):
    doc.add_page_break()


def add_horizontal_rule(doc):
    """Visual divider."""
    p = doc.add_paragraph("_" * 90)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(7)


# =============================================================================
#  PART 1 -- mbXPro_documentation.docx (technical reference)
# =============================================================================

# Bibliography reused across many script chapters
BIB = {
    "mbX": ("Lamichhane U., Lourenco J. (2025). mbX: An R Package for "
            "Streamlined Microbiome Analysis. Stats 8(2):44.",
            "10.3390/stats8020044"),
    "qiime2": ("Bolyen E. et al. (2019). Reproducible, interactive, scalable "
               "and extensible microbiome data science using QIIME 2. "
               "Nature Biotechnology 37:852-857.",
               "10.1038/s41587-019-0209-9"),
    "dada2": ("Callahan B.J. et al. (2016). DADA2: High-resolution sample "
              "inference from Illumina amplicon data. Nature Methods 13:581-583.",
              "10.1038/nmeth.3869"),
    "gg2": ("McDonald D. et al. (2024). Greengenes2 unifies microbial data in "
            "a single reference tree. Nature Biotechnology 42:715-718.",
            "10.1038/s41587-023-01845-1"),
    "cutadapt": ("Martin M. (2011). Cutadapt removes adapter sequences from "
                 "high-throughput sequencing reads. EMBnet.journal 17(1):10-12.",
                 "10.14806/ej.17.1.200"),
    "mafft": ("Katoh K., Standley D.M. (2013). MAFFT multiple sequence "
              "alignment software version 7. MBE 30(4):772-780.",
              "10.1093/molbev/mst010"),
    "fasttree": ("Price M.N., Dehal P.S., Arkin A.P. (2010). FastTree 2 -- "
                 "approximately maximum-likelihood trees for large alignments. "
                 "PLoS ONE 5(3):e9490.",
                 "10.1371/journal.pone.0009490"),
    "hurlbert": ("Hurlbert S.H. (1971). The nonconcept of species diversity: "
                 "a critique and alternative parameters. Ecology 52(4):577-586.",
                 "10.2307/1934145"),
    "chao_shen": ("Chao A., Shen T.J. (2003). Nonparametric estimation of "
                  "Shannon's index of diversity when there are unseen species "
                  "in sample. Environmental and Ecological Statistics 10:429-443.",
                  "10.1023/A:1026096204727"),
    "faith": ("Faith D.P. (1992). Conservation evaluation and phylogenetic "
              "diversity. Biological Conservation 61(1):1-10.",
              "10.1016/0006-3207(92)91201-3"),
    "shannon": ("Shannon C.E. (1948). A mathematical theory of communication. "
                "Bell System Technical Journal 27:379-423.",
                "10.1002/j.1538-7305.1948.tb01338.x"),
    "simpson": ("Simpson E.H. (1949). Measurement of diversity. Nature 163:688.",
                "10.1038/163688a0"),
    "pielou": ("Pielou E.C. (1966). The measurement of diversity in different "
               "types of biological collections. Journal of Theoretical Biology "
               "13:131-144.",
               "10.1016/0022-5193(66)90013-0"),
    "unifrac": ("Lozupone C., Knight R. (2005). UniFrac: a new phylogenetic "
                "method for comparing microbial communities. AEM 71(12):8228-8235.",
                "10.1128/AEM.71.12.8228-8235.2005"),
    "permanova": ("Anderson M.J. (2001). A new method for non-parametric "
                  "multivariate analysis of variance. Austral Ecology 26:32-46.",
                  "10.1111/j.1442-9993.2001.01070.pp.x"),
    "permdisp": ("Anderson M.J. (2006). Distance-based tests for homogeneity "
                 "of multivariate dispersions. Biometrics 62(1):245-253.",
                 "10.1111/j.1541-0420.2005.00440.x"),
    "ancombc2": ("Lin H., Peddada S.D. (2024). Multi-group analysis of "
                 "compositions of microbiomes with covariate adjustments and "
                 "repeated measures. Nature Methods 21:83-91.",
                 "10.1038/s41592-023-02092-7"),
    "picrust2": ("Douglas G.M. et al. (2020). PICRUSt2 for prediction of "
                 "metagenome functions. Nature Biotechnology 38:685-688.",
                 "10.1038/s41587-020-0548-6"),
    "rf": ("Breiman L. (2001). Random Forests. Machine Learning 45(1):5-32.",
           "10.1023/A:1010933404324"),
    "ranger": ("Wright M.N., Ziegler A. (2017). ranger: A Fast Implementation "
               "of Random Forests for High Dimensional Data in C++ and R. "
               "Journal of Statistical Software 77(1):1-17.",
               "10.18637/jss.v77.i01"),
    "proc": ("Robin X. et al. (2011). pROC: an open-source package for R and "
             "S+ to analyze and compare ROC curves. BMC Bioinformatics 12:77.",
             "10.1186/1471-2105-12-77"),
    "aitchison": ("Aitchison J. (1982). The statistical analysis of "
                  "compositional data. JRSS-B 44(2):139-177.",
                  "10.1111/j.2517-6161.1982.tb01195.x"),
    "spearman": ("Spearman C. (1904). The proof and measurement of "
                 "association between two things. Am. J. Psychology 15:72-101.",
                 "10.2307/1412159"),
    "louvain": ("Blondel V.D. et al. (2008). Fast unfolding of communities in "
                "large networks. JSTAT P10008.",
                "10.1088/1742-5468/2008/10/P10008"),
    "igraph": ("Csardi G., Nepusz T. (2006). The igraph software package for "
               "complex network research. InterJournal Complex Systems 1695.",
               "(no DOI -- software)"),
    "psych": ("Revelle W. (2024). psych: Procedures for Psychological, "
              "Psychometric, and Personality Research. R package.",
              "(CRAN: https://cran.r-project.org/package=psych)"),
    "kw": ("Kruskal W.H., Wallis W.A. (1952). Use of ranks in one-criterion "
           "variance analysis. JASA 47(260):583-621.",
           "10.1080/01621459.1952.10483441"),
    "dunn": ("Dunn O.J. (1964). Multiple Comparisons Using Rank Sums. "
             "Technometrics 6(3):241-252.",
             "10.1080/00401706.1964.10490181"),
    "bh": ("Benjamini Y., Hochberg Y. (1995). Controlling the false discovery "
           "rate: a practical and powerful approach to multiple testing. "
           "JRSS-B 57(1):289-300.",
           "10.1111/j.2517-6161.1995.tb02031.x"),
    "vegan": ("Oksanen J. et al. (2022). vegan: Community Ecology Package. "
              "R package version 2.6-4.",
              "(CRAN: https://cran.r-project.org/package=vegan)"),
}


# Each chapter is a dict with all the structured fields the doc needs.
# This data block is intentionally the SINGLE source of truth so future edits
# are all in one place.

CHAPTERS = [
    # ---------------- Step 0 ----------------
    dict(
        id=0,
        script="mbx_primer_identifier.sh",
        title="Auto-detect 16S rRNA primers from raw FASTQ files",
        purpose=(
            "Reads a sample of paired-end FASTQ records and identifies the "
            "exact forward (5'->3') and reverse (5'->3') primers that were "
            "used during library preparation. The primer sequences are needed "
            "by the DADA2 trim-left step (step 4) so that primer bases are "
            "not mistakenly counted as biological sequence."
        ),
        algorithm=(
            "Three-tier detection: "
            "TIER 1 -- IUPAC-aware sliding-window scan (Cutadapt-inspired) "
            "with three orientation hypotheses (RC / DIRECT / SWAP) tested in "
            "parallel. Each candidate is scored across offset windows up to "
            "25 bp (handles Fluidigm CS1/CS2 linkers and frameshifting Ns), "
            "allowing up to 3 mismatches, and the highest-agreement primer "
            "wins. N is treated as wild on BOTH read and primer sides per "
            "canonical IUPAC convention. The primer database is sanitised at "
            "load time (NFKC normalisation + Cyrillic->Latin homoglyph "
            "mapping + strict ASCII IUPAC enforcement). "
            "TIER 2 -- when TIER 1 finds no primer above the rate threshold "
            "(common when the sequencing facility trims primers before "
            "delivery), reads are scanned for conserved 16S V-region anchor "
            "motifs (e.g. V4 starts with TACG.AGG immediately after the 515F "
            "landing site). On a positive hit DETECTION_STATUS is set to "
            "TRIMMED and INFERRED_REGION is filled in. "
            "TIER 3 -- on full failure the script writes rich diagnostic "
            "information into mbx_primer_info.txt: read-length distribution, "
            "per-position A/C/G/T/N composition, top-3 most frequent 5'-"
            "prefixes, and concrete next-step suggestions. Downstream steps "
            "consume DETECTION_STATUS to decide whether to extract the "
            "V-region (DETECTED / USER_SUPPLIED) or skip extract-reads and "
            "train Naive-Bayes on the full Greengenes2 backbone (TRIMMED / "
            "UNKNOWN)."
        ),
        scratch=(
            "Written from scratch by the maintainers. The IUPAC-aware "
            "sliding-window pattern was inspired by the design of Cutadapt; "
            "no source code was reused -- only the algorithmic concept."
        ),
        rules=[
            "Bash 3.2 compatible (no mapfile, no declare -A, no ${var^^}).",
            "PID-based temp filenames (avoids the macOS /tmp XXXXXX bug).",
            "Idempotent: rerunning produces identical output if inputs are unchanged.",
            "Output is placed alongside the FASTQ directory (sibling of), never inside it.",
            "Honors the MBX_OUT_DIR environment variable when set by the orchestrator.",
            "Manual override flags --forward-primer / --reverse-primer skip detection entirely.",
            "Three-tier algorithm: direct match -> V-region anchor -> rich-diagnostic failure.",
            "DETECTION_STATUS is exposed downstream so step 5/6 can auto-fall-back to a full-length classifier when primers are absent.",
        ],
        citations=["cutadapt"],
        inputs=[
            ("<fastq_dir>", "Positional argument. Absolute path to a directory "
                            "containing R1/R2 *.fastq.gz files. Searches the "
                            "given dir + parent + one level of children."),
            ("--samples N",  "Number of reads to sample per file (default 10000)."),
            ("--mismatches K", "Maximum mismatches allowed in a primer match (default 2)."),
            ("--min-match-rate F", "Minimum fraction of sampled reads that must "
                                   "agree on the consensus k-mer (default 0.10)."),
            ("--forward-primer SEQ", "Manual override -- skips detection."),
            ("--reverse-primer SEQ", "Manual override -- skips detection."),
        ],
        upstream_inputs="None. This is the first step in the pipeline.",
        outputs=[
            ("0_primer_handling/mbx_primer_info.txt",
             "Key=value file containing the detected primer sequences, primer "
             "length range, sampled read count, orientation hypothesis used, "
             "per-read agreement rate, plus the new v2 fields "
             "DETECTION_STATUS (DETECTED / TRIMMED / UNKNOWN / USER_SUPPLIED), "
             "CONFIDENCE_LEVEL (HIGH / MEDIUM / LOW / NONE / USER), "
             "INFERRED_REGION (V3-V4, V4, ...), and DETECTION_NOTE."),
        ],
        caveats=[
            "Sequencing facilities sometimes strip primers BEFORE delivery. "
            "v2 detects this automatically: DETECTION_STATUS=TRIMMED is set "
            "and downstream classifier preparation (step 5) auto-falls-back to "
            "training Naive-Bayes on the full Greengenes2 backbone -- the "
            "pipeline continues without user intervention.",
            "When DETECTION_STATUS=UNKNOWN (no primer match AND no V-region "
            "anchor match), the script writes diagnostic information to the "
            "info file: read-length distribution, per-position base "
            "composition, and top-3 most frequent 5'-prefixes. The pipeline "
            "still proceeds (using full-length classifier mode) but the user "
            "should review the diagnostics to confirm the data is 16S.",
            "The detector cannot distinguish biological poly-A/T tails from "
            "primer adapter sequences in pathologically degraded libraries.",
            "Sample size of 10000 reads per file is a heuristic; very low "
            "diversity samples may need --samples 50000 or higher.",
        ],
    ),
    # ---------------- Step 1 ----------------
    dict(
        id=1,
        script="create_manifest.sh",
        title="Build a QIIME2-format sample manifest",
        purpose=(
            "Discovers paired R1/R2 FASTQ files inside the user-supplied FASTQ "
            "directory and produces a tab-separated manifest file in the "
            "format that QIIME2 expects for the import step."
        ),
        algorithm=(
            "Walks the FASTQ directory plus its parent and one level of child "
            "directories, picking up files matching common Illumina naming "
            "conventions (S<N>_L<N>_R{1,2}_<N>.fastq.gz, or simpler variants). "
            "Sample IDs are extracted with two regex passes: first a strict "
            "'s(ample)?[-_]?[0-9]+' match, then a permissive filename-cleanup "
            "fallback. R1 and R2 paths are paired by sample ID using a "
            "Bash-3.2-compatible file-based key/value store (each pseudo-key "
            "is a file under a temp work-dir; this avoids associative arrays)."
        ),
        scratch=(
            "Written from scratch. The manifest format follows the QIIME2 "
            "PairedEndFastqManifestPhred33V2 specification."
        ),
        rules=[
            "Bash 3.2 compatible.",
            "Single-pass discovery; runtime is O(number of files).",
            "Honors the MBX_OUT_DIR environment variable.",
            "Output is placed alongside the FASTQ directory.",
            "Manifest sample IDs match exactly the sample IDs that downstream "
            "steps will look for in the user-supplied metadata.",
        ],
        citations=["qiime2"],
        inputs=[
            ("<fastq_dir>", "Positional. Same as step 0."),
        ],
        upstream_inputs="None analytical -- only the FASTQ dir.",
        outputs=[
            ("1_manifest_file/manifest.txt",
             "Tab-separated file with header 'sample-id', 'forward-absolute-"
             "filepath', 'reverse-absolute-filepath' (or only the forward "
             "column for single-end runs). Read by step 2."),
        ],
        caveats=[
            "Sample IDs derived from filenames must match metadata sample IDs "
            "(this is enforced later in step 4's metadata validator). If the "
            "naming differs, the user must rename FASTQ files OR adjust their "
            "metadata's first column to match.",
            "Single-end / single-direction runs are detected automatically "
            "from the filename pattern; only paired-end was extensively tested.",
        ],
    ),
    # ---------------- Step 2 ----------------
    dict(
        id=2,
        script="artifact_creator.sh",
        title="Import FASTQ files into a QIIME2 .qza artifact",
        purpose=(
            "Wraps `qiime tools import` to convert the manifest plus its "
            "referenced FASTQ files into a single QIIME2 SampleData artifact "
            "(.qza) and produces a quality-summary visualization (.qzv) that "
            "downstream steps will introspect for read length and quality."
        ),
        algorithm=(
            "Auto-detects paired vs single-end from the manifest header, then "
            "calls the appropriate `qiime tools import --type` invocation:\n"
            "  * Paired:  SampleData[PairedEndSequencesWithQuality]\n"
            "  * Single:  SampleData[SequencesWithQuality]\n"
            "Phred33 V2 format is assumed (the only format Illumina has shipped "
            "since 2011). After import, `qiime demux summarize` produces the "
            "QC visualization."
        ),
        scratch="Thin wrapper around standard QIIME2 commands.",
        rules=[
            "Idempotent: skips if .qza already exists.",
            "Derives MBX_OUT_DIR from the manifest path it receives.",
            "Echoes the exact qiime command used (for reproducibility).",
        ],
        citations=["qiime2"],
        inputs=[
            ("<manifest.txt>", "Positional. Path to the manifest produced by step 1."),
        ],
        upstream_inputs="From step 1 (1_manifest_file/manifest.txt).",
        outputs=[
            ("2_first_artifact_file/Paired_End_artifact.qza",
             "QIIME2 SampleData artifact. Read by step 3."),
            ("2_first_artifact_file/Paired_End_artifact.qzv",
             "Quality summary visualization (.qzv). Read by step 3 to choose "
             "DADA2 truncation lengths from the per-position quality scores."),
        ],
        caveats=[
            "Single-end imports are supported but the rest of the pipeline "
            "(notably step 4's DADA2 invocation) is paired-end-centric; "
            "single-end users may need manual edits in step 4.",
        ],
    ),
    # ---------------- Step 3 ----------------
    dict(
        id=3,
        script="mbx_dada2_parameter_finder.sh",
        title="Pick DADA2 truncation lengths from the demux quality profile",
        purpose=(
            "Extracts per-position quality data from the QIIME2 demux "
            "summary, then writes a parameter file containing the optimal "
            "truncation lengths and trim-left values for the DADA2 denoising "
            "step (step 4)."
        ),
        algorithm=(
            "Heuristic: choose the largest position p such that the 25th "
            "percentile of base-quality scores at position p is >= 25. This "
            "is the QIIME2 community's de-facto standard rule for trunc-len. "
            "Trim-left is decided by DETECTION_STATUS from step 0:\n"
            "  - DETECTED or USER_SUPPLIED -> trim-left = primer length "
            "(forward and reverse handled independently from each primer's "
            "own length).\n"
            "  - TRIMMED (sequencing facility already removed primers) -> "
            "trim-left = 0 for both reads. Setting trim-left=20 here would "
            "DESTROY 20 bp of real biological sequence.\n"
            "  - UNKNOWN or info file missing -> trim-left = 20 for both "
            "reads (defensive default; covers the common Illumina-V3 "
            "low-quality-leading-bases case).\n"
            "  - Manual override: --assume-primer-length N forces both "
            "trim-lefts to N regardless of DETECTION_STATUS."
        ),
        scratch=(
            "Heuristic written from scratch following standard QIIME2 community "
            "practice; no specific paper to cite for the trunc-len rule itself."
        ),
        rules=[
            "Reads the .qza/.qzv pair from step 2.",
            "Honors primer info from 0_primer_handling/mbx_primer_info.txt "
            "(DETECTION_STATUS-driven trim-left selection: TRIMMED->0, "
            "DETECTED/USER_SUPPLIED->primer length, UNKNOWN->20).",
            "Calls the helper create_dada2_parameters_txt.sh to format the output.",
            "All paths in the output file are absolute and machine-portable.",
        ],
        citations=["dada2", "qiime2"],
        inputs=[
            ("<artifact.qza>", "Positional. The .qza produced by step 2."),
            ("--primer-info FILE", "Optional. Override auto-discovery of "
                                    "0_primer_handling/mbx_primer_info.txt."),
            ("--assume-primer-length N", "Optional. Force trim-left-f / "
                                          "trim-left-r to N regardless of "
                                          "DETECTION_STATUS."),
        ],
        upstream_inputs=("From step 2 (Paired_End_artifact.qza/.qzv) and step "
                         "0 (mbx_primer_info.txt)."),
        outputs=[
            ("3_dada2_parameters/dada2_parameters.txt",
             "Plain-text parameter file with --p-trunc-len-f, --p-trunc-len-r, "
             "--p-trim-left-f, --p-trim-left-r values; read directly by step 4."),
        ],
        caveats=[
            "If the data are unusually short or low-quality, the 25th-"
            "percentile-25 rule may pick a trunc-len that drops too many reads "
            "in step 4. Users may edit dada2_parameters.txt manually before "
            "step 4 runs.",
        ],
    ),
    # ---------------- Step 4 ----------------
    dict(
        id=4,
        script="mbx_dada2_run.sh",
        title="Run DADA2 denoising and ASV inference",
        purpose=(
            "Wraps `qiime dada2 denoise-paired` to produce the feature table, "
            "representative sequences, and per-sample stats artifact. Includes "
            "a metadata validator that mirrors the validation logic from the "
            "mbX R package so DOWNSTREAM failures are caught at the earliest "
            "possible step."
        ),
        algorithm=(
            "DADA2 (Callahan et al. 2016) is a parametric error-rate model "
            "that infers exact amplicon sequence variants (ASVs) instead of "
            "OTUs. Reads are dereplicated, an error model is fit per-sample, "
            "ASVs are inferred, paired ends are merged, and chimeras are "
            "removed via consensus. CPU count is auto-detected. Memory usage "
            "scales linearly with the number of unique sequences."
        ),
        scratch="Wrapper + a from-scratch Bash 3.2 metadata validator.",
        rules=[
            "Validates metadata BEFORE invoking DADA2 (catches typos and "
            "duplicate sample IDs that would otherwise crash hours into the run).",
            "First metadata column header must be one of: sample-id, id, "
            "sampleid, sample id, featureid, feature id, feature-id "
            "(case-insensitive).",
            "QIIME2 #q2:types directive row is detected and skipped.",
            "Per-sample read counts and chimera rates are reported in step 4's stdout.",
            "Idempotent: skips if all 6 output files exist.",
        ],
        citations=["dada2", "qiime2"],
        inputs=[
            ("<dada2_parameters.txt>", "From step 3."),
            ("<metadata.txt>",          "User-supplied path."),
        ],
        upstream_inputs=("From step 2 (.qza), step 3 (parameter file), "
                         "and the user-supplied metadata file."),
        outputs=[
            ("4_dada2_outputs/feature_table.qza", "ASV count table. Used everywhere downstream."),
            ("4_dada2_outputs/representative_sequences.qza", "ASV sequences. Used by step 5/6."),
            ("4_dada2_outputs/dada2_stats.qza", "Per-sample read filtering report."),
            ("4_dada2_outputs/*.qzv", "Three QC visualizations of the above."),
        ],
        caveats=[
            "DADA2 is the slowest step in the entire pipeline (often 30-60 min "
            "for a typical study). It is single-pass-uninterruptible -- if you "
            "kill it, you must re-run from scratch.",
            "Memory use can spike to several GB for diverse, deep samples.",
            "Chimera removal is consensus-based; rare biological variants may "
            "be incorrectly classified as chimeras in extremely small studies.",
        ],
    ),
    # ---------------- Step 5 ----------------
    dict(
        id=5,
        script="mbx_classifier_arranger.sh",
        title="Download Greengenes2 + prepare the Naive-Bayes classifier",
        purpose=(
            "Downloads the appropriate Greengenes2 release (matched to the "
            "user's QIIME2 version) and writes a manifest describing the "
            "exact reads-extraction and classifier-training commands that "
            "step 6 will execute."
        ),
        algorithm=(
            "Auto-selects GG2 2024.09 for QIIME2 versions >= 2024.5 and GG2 "
            "2022.10 for older releases (the older classifier expects an older "
            "scikit-learn pickle format that 2024+ ships with). Uses `wget -c` "
            "or `curl --continue-at -` for resumable downloads. Computes the "
            "min/max ASV lengths from the exported representative sequences "
            "and writes them into the manifest so step 6 can pass --p-trunc-len "
            "/ --p-min-length / --p-max-length to qiime feature-classifier "
            "extract-reads. "
            "Decides CLASSIFIER_MODE automatically: when both forward and "
            "reverse primers are available, mode='region-specific' (extract-reads "
            "+ V-region NB training); when primers are absent (already trimmed by "
            "the sequencing facility, or detection failed), mode='full-length' "
            "(extract-reads is skipped and step 6 trains NB on the entire GG2 "
            "backbone instead). "
            "NEW in mbX Pro v1.2.0: when CLASSIFIER_MODE=full-length, the "
            "script first attempts to download a pre-trained, sha256-verified "
            "Naive-Bayes classifier from https://zenodo.org/records/20021035 "
            "matching the active QIIME2 version (8 versions are pre-built). "
            "If a compatible classifier is found and verifies, "
            "CLASSIFIER_SOURCE=zenodo is recorded in the manifest, the GG2 "
            "backbone download is skipped, and step 6 will skip both "
            "extract-reads and fit-classifier (saving 30-90 minutes plus "
            "~2 GB of bandwidth). If anything fails (no compatible release, "
            "network down, sha256 mismatch, wrong artifact type), "
            "CLASSIFIER_SOURCE=local-training is recorded and the script "
            "falls back to the pre-1.2 GG2-download + train-locally flow. "
            "The pipeline never aborts because of a Zenodo problem."
        ),
        scratch=(
            "Wrapper logic, GG2-version auto-selection rule, the "
            "CLASSIFIER_MODE auto-fallback (region-specific vs full-length), "
            "and the v1.2.0 Zenodo classifier registry + sha256-verified "
            "fallback chain were all written from scratch."
        ),
        rules=[
            "Resumable downloads (wget -c / curl --continue-at -).",
            "Skips download if .qza files already present (idempotent).",
            "Writes a 'commands ready to run' manifest for step 6 to consume.",
            "CLASSIFIER_MODE auto-fallback: when primers are missing the pipeline "
            "switches to full-length classifier mode and continues without user "
            "intervention.",
            "CLASSIFIER_SOURCE auto-fallback (NEW in v1.2.0): pre-trained "
            "Zenodo classifier first, local training as fallback. Never aborts "
            "because of network or sha256 failures.",
        ],
        citations=["gg2", "qiime2"],
        inputs=[
            ("<mbX_pro_outputs_dir>", "Positional. The shared output directory."),
        ],
        upstream_inputs="From step 4 (representative_sequences.qza).",
        outputs=[
            ("5_classifier_working_dir/<gg2-release>.backbone.full-length.fna.qza", "GG2 reference seqs."),
            ("5_classifier_working_dir/<gg2-release>.backbone.tax.qza", "GG2 reference taxonomy."),
            ("5_classifier_working_dir/length_summary.txt", "ASV length range used for extraction."),
            ("5_classifier_working_dir/mbx_classifier_run_info.txt",
             "Manifest of all paths + the exact commands that step 6 will run."),
        ],
        caveats=[
            "First-time downloads require ~3.5 GB free disk and a stable connection.",
            "Network failures during download are recoverable: re-run and the "
            "download resumes from the last byte received.",
        ],
    ),
    # ---------------- Step 6 ----------------
    dict(
        id=6,
        script="mbx_classifier_run.sh",
        title="Train the Naive-Bayes classifier and assign taxonomy",
        purpose=(
            "Executes the QIIME2 classification workflow described by the "
            "manifest from step 5. Two orthogonal selectors are honoured: "
            "CLASSIFIER_MODE (region-specific vs full-length, decided by the "
            "presence/absence of primers) and CLASSIFIER_SOURCE "
            "(zenodo / cached / local-training / local-training-fallback, "
            "decided by what step 5 was able to obtain)."
        ),
        algorithm=(
            "CLASSIFIER_SOURCE=zenodo or cached: extract-reads AND "
            "fit-classifier-naive-bayes are BOTH skipped; the script "
            "jumps straight to classify-sklearn against the pre-trained "
            "classifier downloaded by step 5 (30-90 minutes saved). "
            "If classify-sklearn errors out (rare sklearn pickle "
            "incompatibility), the script deletes the bad classifier, "
            "calls `mbx_classifier_arranger.sh --skip-zenodo` to ensure "
            "GG2 references are present, retrains the classifier locally, "
            "rewrites CLASSIFIER_SOURCE=local-training-fallback into the "
            "manifest, and retries -- the pipeline never aborts on a "
            "Zenodo problem.\n\n"
            "CLASSIFIER_SOURCE=local-training, region-specific mode "
            "(primers known): "
            "(1) Trim the GG2 full-length references to the user's primer "
            "pair using qiime feature-classifier extract-reads + the "
            "computed min/max ASV length range. "
            "(2) Train a multinomial Naive-Bayes classifier on the trimmed "
            "references and their taxonomic labels. "
            "(3) Apply the trained classifier to the user's representative "
            "sequences.\n\n"
            "CLASSIFIER_SOURCE=local-training, full-length mode (primers "
            "absent and no Zenodo match): "
            "(1) extract-reads is SKIPPED. "
            "(2) Train Naive-Bayes directly on the entire GG2 backbone -- "
            "this takes 30-90 minutes vs 20-60 in region-specific mode and "
            "produces a ~2-3x larger classifier.qza, but is fully automatic "
            "and the trained classifier is reusable for ANY primer set. "
            "(3) classify-sklearn proceeds as usual.\n\n"
            "All paths cache the trained or downloaded classifier on disk so "
            "re-running is near-instant."
        ),
        scratch=(
            "Wrappers around standard QIIME2 commands. The reusable-classifier "
            "caching logic, the dual-mode (region-specific / full-length) "
            "auto-fallback, and the v1.2.0 dual-source "
            "(zenodo / local-training / local-training-fallback) handler "
            "were written from scratch so users whose sequencing facility "
            "trimmed primers, or who hit a Zenodo/sklearn incompatibility, "
            "still complete the pipeline without intervention."
        ),
        rules=[
            "Skips already-completed sub-steps.",
            "Honours CLASSIFIER_MODE from the manifest (region-specific or full-length).",
            "Honours CLASSIFIER_SOURCE from the manifest (zenodo / cached / "
            "local-training / local-training-fallback).",
            "On classify-sklearn failure with a Zenodo classifier, deletes "
            "the bad classifier, falls back to local training, and retries.",
            "Optional --classifier flag lets the user supply a pre-trained classifier (skipping extract + train).",
        ],
        citations=["gg2", "qiime2"],
        inputs=[
            ("<mbX_pro_outputs_dir>", "Positional."),
            ("--classifier <file.qza>", "Optional. Skip steps 5 + 6 training and use a pre-built classifier."),
        ],
        upstream_inputs="From step 5 (manifest + GG2 references) and step 4 (rep seqs).",
        outputs=[
            ("6_classifier_taxonomy/taxonomy.qza", "Per-ASV taxonomic assignment."),
            ("6_classifier_taxonomy/taxonomy.qzv", "QC visualization of the assignments."),
        ],
        caveats=[
            "Naive-Bayes classification depends on training-data quality; for "
            "non-bacterial 16S targets users may want to train on a custom "
            "reference and pass it via --classifier.",
            "Region-specific training takes 20-60 minutes per primer set on a "
            "modern laptop (one-time cost).",
            "Full-length training (auto-selected when primers are absent) takes "
            "30-90 minutes and uses 8-16 GB RAM, but is reusable across ANY "
            "primer set / V-region.",
        ],
    ),
    # ---------------- Step 7 ----------------
    dict(
        id=7,
        script="mbx_taxonomy_run.sh",
        title="Filter mitochondria/chloroplasts and export per-level CSVs",
        purpose=(
            "Removes mitochondrial and chloroplast features (eukaryotic 16S "
            "contaminants common in plant and soil samples) from the feature "
            "table, then exports a CSV file for each of the 7 taxonomic "
            "levels (domain through species)."
        ),
        algorithm=(
            "Calls `qiime taxa filter-table --p-exclude mitochondria,"
            "chloroplast` to drop contaminant features. Then `qiime taxa "
            "barplot` produces a .qzv from which the 7 level-N.csv files "
            "are extracted by exporting the .qzv and reading the "
            "level-<N>.csv that QIIME2 ships inside the visualization."
        ),
        scratch="Wrapper. Filter strategy follows community best practice.",
        rules=[
            "Metadata path is taken as a CLI argument (NOT auto-detected -- "
            "must be supplied explicitly to avoid the wrong file being "
            "picked up).",
            "All 7 levels are exported, even if downstream only uses level 7.",
        ],
        citations=["qiime2"],
        inputs=[
            ("<mbX_pro_outputs_dir>", "Positional."),
            ("<metadata.txt>", "Positional. User-supplied."),
        ],
        upstream_inputs="From step 6 (taxonomy.qza) + step 4 (feature_table.qza).",
        outputs=[
            ("7_taxonomy_csv/feature_table_filtered.qza", "Mito/chloro-removed feature table."),
            ("7_taxonomy_csv/taxa_bar_plots.qzv", "QIIME2 bar-plot visualization."),
            ("7_taxonomy_csv/level-{1..7}.csv", "Tidy-format CSV per taxonomic level."),
            ("7_taxonomy_csv/mbx_taxonomy_info.txt",
             "Provenance file with paths to the above. Read by step 8."),
        ],
        caveats=[
            "Mito/chloro filtering is non-reversible at this point in the "
            "pipeline; if you need them, re-run with the qiime taxa "
            "filter-table command modified.",
            "If the user's metadata file lists samples that are NOT in the "
            "feature table (e.g. because DADA2 dropped them at minimum "
            "frequency), QIIME2 will warn but the export still succeeds.",
        ],
    ),
    # ---------------- Step 8 ----------------
    dict(
        id=8,
        script="mbx_ezclean_all_levels.sh",
        title="mbX::ezclean per taxonomic level",
        purpose=(
            "Calls the mbX R package's ezclean() function on level-7.csv to "
            "consolidate synonymous taxa, drop unassigned features, and write "
            "a tidy wide-format Excel sheet for each of the 7 taxonomic levels."
        ),
        algorithm=(
            "ezclean() parses the full taxonomy string at level 7, splits it "
            "into 7 hierarchical components, then aggregates feature counts "
            "by each level. Identical taxonomic labels (synonymous taxa) are "
            "summed; ambiguous and 'Unassigned' assignments are dropped."
        ),
        scratch=(
            "Wrapper around the mbX::ezclean() R function (Lamichhane & "
            "Lourenco 2025). The wrapper adds: (a) automatic R installation "
            "via Homebrew if Rscript is missing; (b) the _R env-stripping "
            "wrapper to prevent conda-poisoning of R's library path; "
            "(c) per-level error isolation so a species-level failure doesn't "
            "abort the genus-level run."
        ),
        rules=[
            "Uses level-7.csv for ALL 7 levels (level-7 contains the full "
            "taxonomy string, which ezclean parses internally).",
            "setwd() to 8_cleaned_files/ before each ezclean call (the "
            "function writes intermediates to cwd).",
            "Each level runs in its own Rscript invocation so a single "
            "level's failure does not propagate.",
            "_R wrapper unsets R_LIBS, R_HOME, R_LIBS_USER, R_LIBS_SITE etc. "
            "to force the SYSTEM Rscript instead of the conda-shadowed one.",
        ],
        citations=["mbX"],
        inputs=[
            ("<mbX_pro_outputs_dir>", "Positional."),
        ],
        upstream_inputs="From step 7 (level-7.csv).",
        outputs=[
            ("8_cleaned_files/mbX_cleaned_<level>_level-7/mbX_cleaned_<level>_level-7.xlsx",
             "One Excel file per level (domain, phylum, class, order, family, genus, species)."),
        ],
        caveats=[
            "Requires R installed system-wide (Homebrew on macOS). The wrapper "
            "WILL refuse to use a conda-shipped Rscript because that would "
            "trigger Rcpp/methods.dylib loading conflicts.",
            "Excel output is .xlsx (openxlsx); requires no Microsoft software.",
        ],
    ),
    # ---------------- Step 9 ----------------
    dict(
        id=9,
        script="mbx_ezviz_all_levels_all_treatments.sh",
        title="Stacked-bar visualizations (mbX::ezviz)",
        purpose=(
            "Auto-detects every categorical metadata column and produces a "
            "publication-quality stacked-bar PNG for each (taxonomic level x "
            "categorical variable) combination."
        ),
        algorithm=(
            "Categorical columns are detected by an R helper: a column is "
            "categorical if it is not numeric, not all-unique (which would "
            "make it an ID), not constant, and has at least one group with "
            ">= 2 samples. The QIIME2 #q2:types directive row is detected and "
            "removed before column-type inference. ezviz() returns a ggplot "
            "object that the wrapper saves with ggsave at 300 DPI."
        ),
        scratch=(
            "Wrapper around mbX::ezviz(). Categorical-column auto-detection "
            "logic and the directory-name sanitizer were written from scratch."
        ),
        rules=[
            "Bash 3.2-compatible categorical detection via temp R script + sentinels.",
            "Spaces and special characters in metadata column names are "
            "sanitized to underscores when used as directory names.",
            "Idempotent: skips if PNG already exists.",
            "ggsave() is called explicitly because ezviz() does NOT auto-save.",
        ],
        citations=["mbX"],
        inputs=[
            ("<mbX_pro_outputs_dir>", "Positional."),
        ],
        upstream_inputs="From step 8 (cleaned XLSX files) + user metadata.",
        outputs=[
            ("9_visualization_entire/<Variable>/ezviz_<level>_<Variable>.png",
             "One PNG per (level, variable) combination."),
            ("9_visualization_entire/mbx_ezviz_info.txt", "Provenance."),
        ],
        caveats=[
            "Variables with > 30 groups produce visually crowded plots; the "
            "ezviz() default colour palette tops out at 30 distinct colours.",
        ],
    ),
    # ---------------- Step 10 ----------------
    dict(
        id=10,
        script="mbx_ezstat_all_levels_all_treatments.sh",
        title="Kruskal-Wallis + Dunn's post-hoc statistics (mbX::ezstat)",
        purpose=(
            "Tests every taxon at every taxonomic level against every "
            "categorical metadata variable using non-parametric Kruskal-Wallis "
            "and Dunn's post-hoc tests."
        ),
        algorithm=(
            "Kruskal-Wallis (KW; Kruskal & Wallis 1952) is a non-parametric "
            "alternative to ANOVA that does not assume normal-distributed "
            "abundances. P-values are adjusted with Benjamini-Hochberg's FDR "
            "(BH; Benjamini & Hochberg 1995). For variables with > 2 groups, "
            "Dunn's post-hoc pairwise rank tests (Dunn 1964) identify which "
            "specific group pairs differ. Compact-letter-display assignments "
            "(traditional letter codes) are produced for the boxplots."
        ),
        scratch=(
            "Wrapper around mbX::ezstat(). Same categorical-detection helper "
            "as ezviz."
        ),
        rules=[
            "ezstat() handles all xlsx + boxplot writing internally.",
            "Skip-check uses the KW xlsx file (the first file ezstat writes).",
        ],
        citations=["mbX", "kw", "dunn", "bh"],
        inputs=[
            ("<mbX_pro_outputs_dir>", "Positional."),
        ],
        upstream_inputs="From step 8 (cleaned files) + metadata.",
        outputs=[
            ("10_stats/<Variable>/ezstat_KW_<level>_by_<Variable>.xlsx", "Per-taxon KW p + q values."),
            ("10_stats/<Variable>/ezstat_pairwise_<level>_by_<Variable>.xlsx", "Dunn's pairwise."),
            ("10_stats/<Variable>/ezstat_CLD_Summary_<level>_by_<Variable>.xlsx", "CLD letters."),
            ("10_stats/<Variable>/Boxplots_<Variable>/", "Per-taxon PNG boxplots."),
        ],
        caveats=[
            "Kruskal-Wallis assumes independent samples; longitudinal designs "
            "should use ANCOMBC2 (step 14) with random-effects covariates instead.",
        ],
    ),
    # ---------------- Step 11 ----------------
    dict(
        id=11,
        script="mbx_pre_diversity_parameters.sh",
        title="Phylogenetic tree + sampling-depth selection (pre-diversity)",
        purpose=(
            "Runs every analysis that the diversity-significance step (step "
            "12-13) needs as upstream prerequisites: a rooted phylogenetic "
            "tree, per-sample frequency tabulation, and an automatic sampling-"
            "depth recommendation that maximises samples-retained subject to "
            "user-tunable thresholds."
        ),
        algorithm=(
            "(1) MAFFT (Katoh & Standley 2013) aligns the representative "
            "sequences. (2) FastTree (Price et al. 2010) builds the tree. "
            "(3) Sampling-depth selection is group-aware: find the highest "
            "depth where >= 90% of all samples pass AND >= 80% per categorical "
            "group; fall back to >=75% overall, then to Q1. (4) Analytical "
            "rarefaction curves (NOT QIIME2's Monte-Carlo rarefaction) are "
            "computed using the Hurlbert (1971) expected-richness formula and "
            "the Chao-Shen (2003) coverage-adjusted Shannon estimator. "
            "Computation is parallelised across samples via Python's "
            "ProcessPoolExecutor and uses scipy.special.gammaln for numerical "
            "stability. Replacing Monte-Carlo rarefaction with this analytical "
            "formula reduces the step's runtime from hours to seconds with "
            "essentially identical results."
        ),
        scratch=(
            "Original work: the analytical-rarefaction implementation is "
            "novel to mbX Pro (the same approach is used by R packages "
            "iNEXT and vegan::rarecurve, but those are not embedded in any "
            "QIIME2 plugin)."
        ),
        rules=[
            "Skips QIIME2's slow alpha-rarefaction in favour of the analytical version.",
            "Filtered feature table (mito/chloro removed) is preferred when available.",
            "QIIME2 sample-frequencies TSV format quirks (#q2:types row + "
            "comma thousands separators like '28,005.0') are handled in R.",
            "Auto-selects 'most-conservative' categorical column for group "
            "constraint (the column with the most distinct groups).",
        ],
        citations=["mafft", "fasttree", "hurlbert", "chao_shen"],
        inputs=[
            ("<mbX_pro_outputs_dir>", "Positional."),
            ("--skip-tree / --skip-qc", "Reuse a previous run's tree / QC."),
            ("--group-col COL", "Override the auto-selected group constraint column."),
            ("--min-overall F / --min-group F", "Override the depth-selection thresholds."),
        ],
        upstream_inputs=("From step 4 (rep seqs + feature table) and step 7 "
                         "(filtered table + metadata)."),
        outputs=[
            ("11_pre_diversity/rooted-tree.qza", "MAFFT + FastTree rooted tree."),
            ("11_pre_diversity/sample-frequencies.qza", "Per-sample read counts."),
            ("11_pre_diversity/sampling_depth_candidates.csv", "Candidate depths + retention rates."),
            ("11_pre_diversity/alpha_rarefaction_curves.png", "Analytical rarefaction figure."),
            ("11_pre_diversity/mbx_pre_diversity_info.txt", "Provenance read by step 12."),
        ],
        caveats=[
            "The analytical rarefaction assumes Hurlbert's i.i.d. sampling "
            "model; this is the same assumption every other rarefaction "
            "method makes and is generally accepted in the field.",
            "Faith's PD is NOT computed in the rarefaction curve (it is "
            "computed in step 12's core-metrics step regardless; including it "
            "in the curve adds ~70% runtime for no scientific gain because "
            "the tree structure does not change with depth).",
        ],
    ),
    # ---------------- Step 12 ----------------
    dict(
        id=12,
        script="mbx_alpha_diversity_run.sh",
        title="Alpha diversity metrics + per-variable statistics",
        purpose=(
            "Computes the five canonical alpha-diversity metrics (ASV richness, "
            "Shannon, Simpson, Faith's PD, Pielou's evenness) per sample, "
            "tabulates them in a single Excel file, and runs Kruskal-Wallis + "
            "Dunn's tests + boxplots per categorical variable."
        ),
        algorithm=(
            "QIIME2's `qiime diversity alpha-phylogenetic` and `qiime diversity "
            "alpha` are invoked for each metric, the per-sample vectors are "
            "extracted, and the five vectors are joined on sample-id into a "
            "wide table. The categorical-detection helper (same as ezviz / "
            "ezstat) iterates over every categorical variable to produce one "
            "Boxplots/<variable>/ + stats/<variable>/ pair per metric."
        ),
        scratch=(
            "Composition of standard QIIME2 alpha-diversity calls + a "
            "from-scratch tabulator + the same KW + Dunn machinery as ezstat."
        ),
        rules=[
            "Sampling depth is taken from step 11's mbx_pre_diversity_info.txt.",
            "All metrics are computed on the SAME rarefied table for consistency.",
            "Per-sample alpha values are the SAME values used by step 18's "
            "sample-inventory section.",
        ],
        citations=["faith", "shannon", "simpson", "pielou", "kw", "dunn", "bh"],
        inputs=[
            ("<mbX_pro_outputs_dir>", "Positional."),
        ],
        upstream_inputs="From step 11 (rooted tree, depth) + step 7 (filtered table) + metadata.",
        outputs=[
            ("12_alpha_diversity_results/all_alpha_outputs/", "Per-metric .qza files."),
            ("12_alpha_diversity_results/alpha_diversity.xlsx",
             "Wide-table with columns: sample-id, ASVs_or_Features, Shannon_Index, "
             "Simpson_Diversity, Faith_Phylogenetic_Diversity, Pielou_Evenness."),
            ("12_alpha_diversity_results/stats_for_alpha_diversity/<Var>/", "KW + Dunn xlsx."),
            ("12_alpha_diversity_results/bloxplots_for_alpha_diversity/<Var>/", "Per-metric PNG boxplots."),
        ],
        caveats=[
            "Faith's PD requires the rooted tree from step 11; if step 11 "
            "was skipped, only the four non-phylogenetic metrics are produced.",
        ],
    ),
    # ---------------- Step 13 ----------------
    dict(
        id=13,
        script="mbx_beta_diversity_run.sh",
        title="Beta diversity (PCoA, PERMANOVA, PERMDISP, Adonis)",
        purpose=(
            "Computes pairwise sample distance matrices using four metrics, "
            "produces ordination figures coloured by every categorical "
            "variable, runs PERMANOVA significance tests with pairwise "
            "follow-ups, tests for homogeneity of dispersion (PERMDISP), "
            "produces a multivariate Adonis model on all metadata columns "
            "simultaneously, plus a clustered distance heatmap and UPGMA "
            "dendrogram for the all-samples view."
        ),
        algorithm=(
            "Distance metrics: Bray-Curtis (abundance-based), Jaccard "
            "(presence/absence), weighted UniFrac (abundance + phylogeny), "
            "unweighted UniFrac (Lozupone & Knight 2005). Tests: PERMANOVA "
            "(Anderson 2001) for between-group differences, with "
            "Benjamini-Hochberg-adjusted pairwise comparisons; PERMDISP "
            "(Anderson 2006) for dispersion homogeneity (the test that "
            "validates whether PERMANOVA results are due to mean shifts vs "
            "spread differences); Adonis multivariate model for joint "
            "explanation of variance by all metadata columns."
        ),
        scratch=(
            "Composition of standard QIIME2 beta-diversity calls + custom "
            "Adonis multivariate model + custom distance heatmap + UPGMA "
            "tree (all in R)."
        ),
        rules=[
            "Per-categorical-variable subdirectory layout matches the rest "
            "of the pipeline.",
            "Working files (intermediate matrices) live in working_dir_beta_"
            "diversity/ and are NOT considered publication outputs.",
        ],
        citations=["unifrac", "permanova", "permdisp", "vegan"],
        inputs=[
            ("<mbX_pro_outputs_dir>", "Positional."),
        ],
        upstream_inputs="From step 11 (tree + depth) + step 7 (filtered table) + metadata.",
        outputs=[
            ("13_beta_diversity_results/results_by_categorical_variables/<Var>/PCoA_*.png", "Ordinations."),
            ("13_beta_diversity_results/results_by_categorical_variables/<Var>/PERMANOVA_*.xlsx", "Significance."),
            ("13_beta_diversity_results/results_by_categorical_variables/<Var>/Pairwise_*.xlsx", "Pairwise."),
            ("13_beta_diversity_results/results_by_categorical_variables/<Var>/PERMDISP_*.xlsx", "Dispersion."),
            ("13_beta_diversity_results/all_samples_beta_diversity/Adonis_*.xlsx", "Multivariate."),
            ("13_beta_diversity_results/all_samples_beta_diversity/heatmap_*.png", "Distance heatmap."),
            ("13_beta_diversity_results/all_samples_beta_diversity/UPGMA_*.png", "Dendrogram."),
        ],
        caveats=[
            "PERMANOVA with very small per-group sample sizes (n<5) has low "
            "statistical power; users should interpret marginal p-values cautiously.",
            "UniFrac requires the rooted tree; if step 11 was skipped, only "
            "Bray-Curtis and Jaccard are produced.",
        ],
    ),
    # ---------------- Step 14 ----------------
    dict(
        id=14,
        script="mbx_ancombc2_run.sh",
        title="Differential abundance via ANCOMBC2",
        purpose=(
            "Runs the ANCOM-BC2 (Lin & Peddada 2024) compositional-aware "
            "differential abundance test for every taxonomic level x every "
            "categorical metadata variable, with both pairwise and global "
            "(reference-free) testing enabled."
        ),
        algorithm=(
            "ANCOM-BC2 fits a log-linear regression on the bias-corrected "
            "abundance scale, controlling for compositionality. Setting "
            "pairwise=TRUE produces every pairwise contrast within each "
            "variable; global=TRUE adds an omnibus reference-free test. "
            "Multiple-testing correction is Holm (default) plus Benjamini-"
            "Hochberg, with q-value columns named 'q_<contrast>'."
        ),
        scratch=(
            "Wrapper around the ANCOMBC2 R Bioconductor package + a from-"
            "scratch Mac-Apple-Silicon-aware dependency installer that "
            "creates a dedicated R conda env so it never pollutes QIIME2's."
        ),
        rules=[
            "Reference is the alphabetically-first level of each variable "
            "(R's modeling default); pairwise + global are run regardless "
            "so the reference choice affects only labelling, not conclusions.",
            "Working files live in working_dir_differential_abundance/.",
            "Master roll-up Summary_all_levels_all_variables.xlsx aggregates "
            "every (level x variable x contrast) significant taxon.",
        ],
        citations=["ancombc2"],
        inputs=[
            ("<mbX_pro_outputs_dir>", "Positional."),
        ],
        upstream_inputs="From step 7 (filtered table) + metadata.",
        outputs=[
            ("14_differential_abundance_ANCOMBC2/ANCOMBC2_<level>/", "Per-level result dir."),
            (".../ancombc2_primary_<level>_<variable>.xlsx", "Per-(level,variable) DA table."),
            ("14_.../Summary_all_levels_all_variables.xlsx", "Master roll-up read by step 18."),
        ],
        caveats=[
            "ANCOMBC2 requires reasonably-sized groups (n>=4 per group recommended); "
            "smaller groups may yield no significant taxa even when biologically real.",
            "Compositional methods cannot detect 'absolute' abundance changes -- "
            "only relative shifts.",
        ],
    ),
    # ---------------- Step 15 ----------------
    dict(
        id=15,
        script="mbx_picrust_run.sh",
        title="Functional prediction via PICRUSt2",
        purpose=(
            "Predicts functional gene content (KO, EC, MetaCyc pathway, COG) "
            "from the user's 16S rRNA ASVs using the PICRUSt2 pipeline, then "
            "produces stacked-bar plots of the top 20 MetaCyc pathways, "
            "differential-abundance heatmaps, a Spearman-correlation heatmap "
            "linking ANCOMBC2 genera to differentially-abundant pathways, and "
            "a Bray-Curtis functional beta-diversity PCoA."
        ),
        algorithm=(
            "PICRUSt2 (Douglas et al. 2020) places ASVs onto a reference tree, "
            "then maps placements to genome content predictions. NSTI (Nearest "
            "Sequenced Taxon Index) measures placement reliability; ASVs with "
            "NSTI > 2 are dropped (default), with a per-sample warning when "
            "any sample's mean NSTI > 1.0. PICRUSt2 is auto-installed into a "
            "DEDICATED conda env (picrust2-mbx) on first run so it does NOT "
            "contaminate the user's QIIME2 env. "
            "COG predictions are skipped on PICRUSt2 >= 2.6.x because that "
            "version's COG database is incompatible with the predict_metagenomes "
            "API; this is reported as an explicit limitation in the report."
        ),
        scratch=(
            "Wrapper. The conda-env-isolation pattern, the per-sample NSTI "
            "warning system, the genus<->pathway correlation heatmap, and the "
            "functional-beta-diversity PCoA are original to mbX Pro."
        ),
        rules=[
            "Configurable NSTI threshold via --nsti (default 2.0).",
            "Per-categorical-variable subdir layout.",
            "All plots are subtitle-annotated 'Predicted functional profiles "
            "-- inferred from 16S rRNA data' to prevent misinterpretation.",
            "All plots exported at >=300 DPI as both PNG and PDF.",
        ],
        citations=["picrust2", "spearman"],
        inputs=[
            ("<mbX_pro_outputs_dir>", "Positional."),
            ("--nsti F", "NSTI threshold (default 2.0)."),
            ("--variables V1,V2,...", "Limit analysis to specific metadata variables."),
        ],
        upstream_inputs=("From step 4 (rep seqs + feature table) + step 7 "
                         "(filtered table) + step 14 (ANCOMBC2 results, for "
                         "the genus<->pathway correlation step) + metadata."),
        outputs=[
            ("15_picrust2/working_dir_picrust2/", "Intermediate placement + KO files."),
            ("15_picrust2/picrust2_<Variable>/", "Per-variable analyses."),
            ("15_picrust2/picrust2_report.html", "Self-contained PICRUSt2 sub-report."),
            ("15_picrust2/mbx_picrust2_info.txt", "Provenance."),
        ],
        caveats=[
            "PICRUSt2 PREDICTIONS are not measured gene content -- they are "
            "inferences. Treat results as hypotheses for follow-up shotgun work.",
            "First run must download the PICRUSt2 reference database (~3 GB).",
            "COG predictions are skipped on PICRUSt2 >=2.6 due to upstream API "
            "incompatibility (this is documented as a limitation in the report).",
        ],
    ),
    # ---------------- Step 16 ----------------
    dict(
        id=16,
        script="mbx_ml_classifier_run.sh",
        title="Random Forest biomarker classifier",
        purpose=(
            "Trains a Random Forest classifier (Breiman 2001) on each "
            "(taxonomic level x categorical variable) combination, with "
            "automatic cross-validation strategy selection, class-imbalance "
            "handling, full multi-class diagnostics, and SHAP-style local "
            "feature-importance heatmaps."
        ),
        algorithm=(
            "Implementation uses `ranger` (Wright & Ziegler 2017) for speed. "
            "Cross-validation is auto-selected: 5-fold stratified when "
            "n>=20 samples per class, leave-one-out otherwise. Class imbalance "
            "is handled via `case.weights` proportional to inverse class size. "
            "Outputs include accuracy, macro-averaged AUC (one-vs-rest, via "
            "pROC), per-class sensitivity/specificity/F1, permutation feature "
            "importance, and per-sample local importance values rendered as a "
            "SHAP-style heatmap. Confusion matrix + ROC curves are saved per "
            "model. The trained model itself is saved as model.rds for re-use."
        ),
        scratch=(
            "Wrapper + the auto-CV decision rule + the SHAP-style local-"
            "importance heatmap layout were written from scratch."
        ),
        rules=[
            "Per-variable summary xlsx aggregates one row per taxonomic level.",
            "Class-filtering: classes with < 3 members are dropped before training.",
            "Zero-variance feature filter: features that are constant across "
            "kept samples are dropped before training.",
        ],
        citations=["rf", "ranger", "proc"],
        inputs=[
            ("<mbX_pro_outputs_dir>", "Positional."),
            ("--levels g,f,s", "Comma-separated subset of taxonomic levels."),
        ],
        upstream_inputs=("From step 8 (cleaned files for each level) + "
                         "step 7 (metadata)."),
        outputs=[
            ("16_ml_biomarkers/<Variable>/<level>/", "Per-(level, variable) results."),
            (".../confusion_matrix.png .../roc_curves.png .../feature_importance.png", "Diagnostics."),
            (".../shap_local_importance_heatmap.png", "Per-sample local importance."),
            (".../model.rds", "Trained ranger model object."),
            ("16_ml_biomarkers/<Variable>/Summary_RF_<Variable>.xlsx", "Per-variable roll-up."),
        ],
        caveats=[
            "Random Forest variable-importance is a CORRELATIVE signal, not "
            "a causal one; reports flag this explicitly.",
            "AUC requires >=2 classes per fold; LOOCV may produce undefined AUC "
            "for some folds (the script reports macro-averaged AUC ignoring NA folds).",
            "For very small datasets (n<10), even LOOCV is unreliable; the "
            "report prints a 'small-N' warning when n<10.",
        ],
    ),
    # ---------------- Step 17 ----------------
    dict(
        id=17,
        script="mbx_network_run.sh",
        title="CLR + Spearman co-occurrence networks",
        purpose=(
            "Builds compositional-aware co-occurrence networks for the global "
            "samples and for each level of each categorical variable; "
            "identifies hub taxa, Louvain communities, and group-specific "
            "network differences."
        ),
        algorithm=(
            "Pipeline per network: "
            "(1) Prevalence filter (minimum number of samples a taxon must "
            "appear in to be considered). "
            "(2) Pseudocount + centred-log-ratio (CLR) transform (Aitchison "
            "1982) to convert compositional counts to a real-valued space. "
            "(3) Spearman correlation matrix + Benjamini-Hochberg FDR via "
            "psych::corr.test. "
            "(4) Edge filter on |rho| AND q-value. "
            "(5) Build igraph (Csardi & Nepusz 2006) graph; compute degree, "
            "betweenness, hub score (Kleinberg). "
            "(6) Louvain modularity (Blondel et al. 2008) for community "
            "detection. "
            "(7) Hub identification (top decile by hub score). "
            "Per-group networks are compared to the global network via Jaccard "
            "overlap of hub sets and a multi-panel plot."
        ),
        scratch=(
            "Original work: the entire pipeline is written from scratch. "
            "Underlying methods (CLR, Spearman, Louvain, igraph metrics) are "
            "all from the cited literature."
        ),
        rules=[
            "Taxonomic levels can be subsetted with --levels (default g,f,s).",
            "Status flags (NO_EDGES, TOO_SMALL, SKIPPED_FEW_TAXA) are "
            "preserved through the aggregator so the report can document why "
            "a particular subgroup did not produce a network.",
            "Re-runs rebuild the summary lines from on-disk run.log files so "
            "idempotent re-runs do not lose data.",
        ],
        citations=["aitchison", "spearman", "bh", "louvain", "igraph", "psych"],
        inputs=[
            ("<mbX_pro_outputs_dir>", "Positional."),
            ("--levels g,f,s", "Subset of taxonomic levels."),
            ("--min-prev N",   "Minimum prevalence (samples) per taxon."),
            ("--rho-min F",    "Minimum |rho| for an edge (default 0.6)."),
            ("--q-max F",      "Maximum BH q-value for an edge (default 0.05)."),
        ],
        upstream_inputs="From step 8 (cleaned files) + metadata.",
        outputs=[
            ("17_co_occurrence_networks/global_<level>/edges.tsv .../nodes.tsv .../network.graphml",
             "Tabular + GraphML for downstream tools (Cytoscape, Gephi)."),
            ("17_co_occurrence_networks/global_<level>/network_plot.{png,pdf}", "Embedded in report."),
            ("17_co_occurrence_networks/global_<level>/network_summary.xlsx", "Network metrics."),
            ("17_co_occurrence_networks/global_<level>/hub_taxa.xlsx", "Hub taxa list."),
            ("17_co_occurrence_networks/per_var_<Variable>/<level>/<group>/", "Per-group networks."),
            ("17_co_occurrence_networks/per_var_<Variable>/group_comparison.xlsx", "Cross-group metrics + Jaccard."),
            ("17_co_occurrence_networks/mbx_networks_info.txt", "Provenance."),
        ],
        caveats=[
            "GROUPS is a Bash built-in array of user group IDs; we renamed our "
            "internal variable to GRP_LIST to avoid silent failures.",
            "Networks with very few taxa (<10 after prevalence filter) are "
            "skipped (status=SKIPPED_FEW_TAXA) and reported in the summary.",
            "Spearman + BH-FDR is the most conservative compositional-aware "
            "edge inference; users wanting more permissive networks can lower "
            "--rho-min to 0.4 or --q-max to 0.10.",
        ],
    ),
    # ---------------- Step 18 ----------------
    dict(
        id=18,
        script="mbx_final_report.sh",
        title="Consolidated HTML + PDF final report",
        purpose=(
            "Aggregates the outputs of every preceding step into a single "
            "self-contained HTML report and a publication-ready A4 PDF."
        ),
        algorithm=(
            "Discovers which steps were run by checking for each step's "
            "sentinel info file. For each step, dedicated R extractor "
            "functions read the sentinel file and any tabular outputs (xlsx, "
            "tsv, csv) and convert them into HTML tables. Figures are base64-"
            "embedded so the HTML is a single file with no external "
            "dependencies. The PDF is rendered via headless Chrome (preferred), "
            "wkhtmltopdf, or macOS cupsfilter (last resort), all driven by "
            "@page A4 CSS rules. A cross-step convergence table identifies "
            "taxa flagged by 2+ of {ANCOMBC2, RF top-importance, network hub} "
            "as biomarker candidates worth follow-up validation."
        ),
        scratch=(
            "Original work. The convergence table, the per-step 'not run' "
            "graceful boxes (with the exact command needed to add the step), "
            "and the print-CSS hardening (single-column figure grid + 84vh "
            "image height limit) are all novel."
        ),
        rules=[
            "Self-contained: HTML stays valid even if the directory is moved.",
            "All paths shown in the HTML are RELATIVE to the output directory; "
            "all paths in *_info.txt files remain ABSOLUTE.",
            "htmltools::save_html() is used (NOT writeLines(as.character()) -- "
            "the latter silently drops <head> due to a tags$head() hoisting quirk).",
            "Logo is auto-discovered from script-dir, ~/bin, /usr/local/share/...",
        ],
        citations=["mbX"],
        inputs=[
            ("<mbX_pro_outputs_dir>", "Positional."),
            ("--no-pdf",        "Skip PDF rendering."),
            ("--pdf-engine X",  "Force chrome / wkhtmltopdf / cupsfilter."),
            ("--force-rerun",   "Regenerate even if 18_final_report exists."),
        ],
        upstream_inputs="From every previous step (sentinel files + outputs).",
        outputs=[
            ("18_final_report/mbX_pro_final_report.html", "Single-file HTML report."),
            ("18_final_report/mbX_pro_final_report.pdf",  "A4 publication-ready PDF."),
            ("18_final_report/mbx_final_report_info.txt", "Provenance."),
        ],
        caveats=[
            "PDF generation requires Chrome (best), wkhtmltopdf (fallback), "
            "or cupsfilter (lowest quality, macOS-only). The HTML is always "
            "produced even without a PDF engine.",
            "Reports for very deep studies (>500 samples, >10000 ASVs) can "
            "exceed 50 MB; this is unavoidable when embedding figures at "
            "publication resolution.",
        ],
    ),
    # ---------------- Step 19 (orchestrator) ----------------
    dict(
        id=19,
        script="mbXPro",
        title="One-command orchestrator",
        purpose=(
            "The single command that the lay user actually runs. Invokes "
            "every preceding step in the correct order using one shared "
            "output directory. The user does not need to know any internal "
            "command, parameter, or file path."
        ),
        algorithm=(
            "(1) Pre-flight: validate FASTQ dir, metadata file, conda + "
            "QIIME2 + R availability, and that all step scripts are in PATH. "
            "(2) Compute exactly ONE timestamp; pre-create "
            "<parent_of_FASTQ>/mbX_pro_outputs_<TIMESTAMP>/. "
            "(3) Export MBX_OUT_DIR so steps 0 and 1 (which previously each "
            "made their own timestamped dirs) reuse the shared directory. "
            "(4) Iterate steps 0-18, calling each script with its correct "
            "argument signature, tee'ing per-step stdout/stderr to a log file. "
            "(5) Per-step idempotency checks let --resume safely skip "
            "already-completed steps. (6) On completion, write provenance + "
            "print a per-step timeline table."
        ),
        scratch="Original work, written from scratch.",
        rules=[
            "Exactly ONE mbX_pro_outputs_* directory is ever created per run.",
            "Bash 3.2 compatible.",
            "MBX_OUT_DIR env var contract is shared with the patched step 0 "
            "and step 1 scripts.",
            "Resume mode finds the most-recent matching dir and continues.",
            "Failure aborts the pipeline by default; --keep-going overrides.",
        ],
        citations=[],
        inputs=[
            ("<fastq_dir>",     "Positional. Directory containing raw FASTQ files."),
            ("<metadata_file>", "Positional. QIIME2-format metadata."),
            ("--resume",        "Reuse the most-recent existing output dir."),
            ("--out DIR",       "Force a specific output directory."),
            ("--from N --to N", "Run only a sub-range of steps."),
            ("--skip A,B,C",    "Skip specific steps by ID."),
            ("--forward-primer SEQ / --reverse-primer SEQ", "Manual primer override."),
            ("--no-pdf",        "Skip PDF rendering in step 18."),
            ("--threads N",     "Override auto-detected CPU count."),
            ("--keep-going",    "Continue past failing steps."),
            ("--dry-run",       "Print commands without executing."),
        ],
        upstream_inputs="None -- this is the entry point.",
        outputs=[
            ("mbX_pro_outputs_<TIMESTAMP>/", "The single shared output dir."),
            ("mbX_pro_outputs_<TIMESTAMP>/mbXPro_run_info.txt",
             "Provenance: software versions, per-step status + timing, user flags."),
            ("mbX_pro_outputs_<TIMESTAMP>/_pipeline_log/", "Per-step + master log files."),
        ],
        caveats=[
            "Requires that the user has activated the QIIME2 conda environment "
            "before invoking; the orchestrator does NOT auto-activate to avoid "
            "fragile conda-init issues.",
            "On HPC, run inside an interactive job allocation or wrap with "
            "your scheduler (sbatch / qsub) to avoid login-node enforcement.",
        ],
    ),
]


def build_technical_doc(out_path):
    doc = Document()
    set_default_font(doc, name="Calibri", size=11)

    # --- Cover page ---
    add_logo_centered(doc, width_in=3.0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("mbX Pro")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Comprehensive Technical Documentation")
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x5A, 0x64, 0x70)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Single-command, end-to-end 16S rRNA microbiome pipeline.\n"
        "From raw FASTQ to publication-ready figures, statistics, and an HTML + PDF report."
    )
    r.italic = True

    doc.add_paragraph("\n\n")
    add_table_2col(doc, [
        ("Version",   "1.0.0"),
        ("Authors",   "Utsav Lamichhane, Jeferson Lourenco"),
        ("Affiliation", "Department of Animal and Dairy Science, University of Georgia"),
        ("License",   "MIT"),
        ("Citation",  "Lamichhane U., Lourenco J. (2025). mbX: An R Package for "
                      "Streamlined Microbiome Analysis. Stats 8(2):44. "
                      "doi:10.3390/stats8020044"),
    ])
    add_page_break(doc)

    # --- Foreword ---
    add_h1(doc, "1. Foreword")
    add_para(doc, (
        "mbX Pro is a deterministic, single-command microbiome analysis "
        "pipeline. The user invokes one command -- mbXPro <fastq_dir> "
        "<metadata_file> -- and the pipeline executes every step from "
        "primer detection through final-report generation in a single "
        "shared output directory."
    ))
    add_para(doc, (
        "This document is the comprehensive TECHNICAL reference. It is "
        "intended for collaborators, peer reviewers, and developers who "
        "need to understand every algorithm, every input/output contract, "
        "and every caveat of every script. It does NOT contain run-time "
        "paths or any user-specific data."
    ))
    add_para(doc, (
        "If you only want to USE the pipeline, see the companion document "
        "how_to_run_mbX_Pro.docx. If you are publishing the pipeline to "
        "GitHub, see how_to_upload_mbX_Pro.docx."
    ))

    # --- Universal design rules ---
    add_h1(doc, "2. Universal design rules")
    add_para(doc, (
        "Every script in the pipeline follows the rules in this section. "
        "Per-script chapters list ADDITIONAL rules specific to that script."
    ))

    add_h2(doc, "2.1 Bash 3.2 compatibility")
    add_para(doc, (
        "macOS still ships with Bash 3.2 by default. To remain portable across "
        "every supported platform, NO script in this pipeline uses any of the "
        "following Bash 4+ features:"
    ))
    for f in ["mapfile / readarray (use a `while IFS= read -r` loop instead)",
              "declare -A (associative arrays; use a temp directory file-based key/value store)",
              "${var^^} / ${var,,} (case conversion; use `tr '[:upper:]' '[:lower:]'`)",
              "realpath (use a `cd && pwd` pattern)"]:
        add_bullet(doc, f)

    add_h2(doc, "2.2 PID-based temp filenames")
    add_para(doc, (
        "macOS Bash 3.2's mktemp leaves a literal 'XXXXXX' suffix file on "
        "disk if the script is killed mid-run. Every temp-file-creating "
        "script uses a PID + epoch suffix instead:"
    ))
    add_code_block(doc,
        '_TMPID="${$}_$(date +%s)"\n'
        'FILE="/tmp/mbx_name_${_TMPID}.R"\n'
        'trap \'rm -f "/tmp/mbx_*_${_TMPID}*" 2>/dev/null\' EXIT')

    add_h2(doc, "2.3 Idempotency")
    add_para(doc, (
        "Every step checks whether its primary output already exists; if so "
        "it logs 'CACHED' and skips. This means the pipeline can be re-run "
        "after a crash with --resume and only the failed step is re-executed."
    ))

    add_h2(doc, "2.4 Friendly errors")
    add_para(doc, (
        "Every err() call includes BOTH the cause AND the fix instruction. "
        "Errors should always be actionable by a non-expert user."
    ))

    add_h2(doc, "2.5 Auto-detect CPU cores")
    add_code_block(doc,
        'if command -v nproc &>/dev/null; then N_JOBS="$(nproc)"\n'
        'elif command -v sysctl &>/dev/null; then N_JOBS="$(sysctl -n hw.logicalcpu 2>/dev/null || echo 1)"\n'
        'else N_JOBS=1; fi')

    add_h2(doc, "2.6 Output directory placement")
    add_para(doc, (
        "Output is placed ALONGSIDE (sibling of) the input FASTQ directory, "
        "never inside it and never in $(pwd):"
    ))
    add_code_block(doc, 'OUT_ROOT="$(dirname "$FASTQ_DIR")/mbX_pro_outputs_${TIMESTAMP}"')
    add_para(doc, (
        "When invoked by the mbXPro orchestrator, the MBX_OUT_DIR environment "
        "variable is set, and the early scripts (steps 0 and 1) reuse it "
        "instead of creating their own timestamped dir. This guarantees that "
        "exactly ONE mbX_pro_outputs_* directory is created per pipeline run."
    ))

    add_h2(doc, "2.7 R scripts written via echo statements")
    add_para(doc, (
        "When a Bash wrapper needs to execute an R script, the R script is "
        "written to a temp file via echo statements (NOT a heredoc). This "
        "avoids quoting conflicts with Python and HTML content embedded in "
        "the R code."
    ))

    add_h2(doc, "2.8 Info file convention")
    add_para(doc, (
        "Every step writes an *_info.txt provenance file in its output "
        "subdirectory using simple key=value lines. Downstream steps read "
        "these files via the helper:"
    ))
    add_code_block(doc, '_read_key() { grep "^${1}=" "$2" 2>/dev/null | cut -d= -f2-; }')

    add_h2(doc, "2.9 R environment isolation")
    add_para(doc, (
        "All R-based steps strip the conda-shipped R environment variables "
        "(R_LIBS, R_HOME, R_LIBS_USER, R_LIBS_SITE, ...) and prefer the "
        "system Rscript. This avoids the well-known Rcpp.so / methods.dylib "
        "loading conflict that occurs when QIIME2's conda env shadows the "
        "system R installation."
    ))

    add_page_break(doc)

    # --- Per-step chapters ---
    add_h1(doc, "3. Pipeline reference")
    add_para(doc, (
        "The 19 chapters below cover every script in the order in which the "
        "orchestrator (mbXPro) executes them."
    ))

    for ch in CHAPTERS:
        add_page_break(doc)
        add_h1(doc, f"3.{ch['id']+1}  Step {ch['id']} -- {ch['script']}")
        add_para(doc, ch["title"], italic=True)

        add_h3(doc, "Purpose")
        add_para(doc, ch["purpose"])

        add_h3(doc, "Algorithm")
        add_para(doc, ch["algorithm"])

        add_h3(doc, "Provenance")
        add_para(doc, ch["scratch"])

        add_h3(doc, "Rules followed (in addition to the universal rules above)")
        for r in ch["rules"]:
            add_bullet(doc, r)

        add_h3(doc, "Citations")
        if ch["citations"]:
            for cid in ch["citations"]:
                full, doi = BIB[cid]
                add_bullet(doc, f"{full}  [DOI: {doi}]")
        else:
            add_para(doc, "(none -- original work)", italic=True)

        add_h3(doc, "Inputs")
        add_para(doc, "Direct command-line arguments:")
        for k, v in ch["inputs"]:
            add_bullet(doc, f"{k} -- {v}")
        add_para(doc, "Where the inputs come from inside the pipeline:")
        add_para(doc, ch["upstream_inputs"], italic=True)

        add_h3(doc, "Outputs")
        for path, desc in ch["outputs"]:
            add_bullet(doc, f"{path} -- {desc}")

        add_h3(doc, "Caveats")
        for c in ch["caveats"]:
            add_bullet(doc, c)

    # --- Pipeline data flow appendix ---
    add_page_break(doc)
    add_h1(doc, "4. Pipeline data flow (appendix)")
    add_para(doc, (
        "This appendix shows, for every step, which earlier step's output is "
        "consumed."
    ))
    flow_rows = [
        ("0  -> ",  "(reads raw FASTQ)"),
        ("1  -> ",  "(reads raw FASTQ)"),
        ("2  -> ",  "1 (manifest)"),
        ("3  -> ",  "2 (artifact + visualisation), 0 (primer info)"),
        ("4  -> ",  "3 (parameters), user metadata"),
        ("5  -> ",  "4 (rep seqs)"),
        ("6  -> ",  "5 (manifest + GG2 refs), 4 (rep seqs)"),
        ("7  -> ",  "6 (taxonomy), 4 (feature table), user metadata"),
        ("8  -> ",  "7 (level-7.csv)"),
        ("9  -> ",  "8 (cleaned files), user metadata"),
        ("10 -> ",  "8 (cleaned files), user metadata"),
        ("11 -> ",  "4 (rep seqs + table), 7 (filtered table), user metadata"),
        ("12 -> ",  "11 (tree + depth), 7 (filtered table), user metadata"),
        ("13 -> ",  "11 (tree + depth), 7 (filtered table), user metadata"),
        ("14 -> ",  "7 (filtered table), user metadata"),
        ("15 -> ",  "4 (rep seqs + table), 7 (filtered), 14 (DA table), user metadata"),
        ("16 -> ",  "8 (cleaned files), user metadata"),
        ("17 -> ",  "8 (cleaned files), user metadata"),
        ("18 -> ",  "every previous step (sentinel files + outputs)"),
    ]
    add_table_2col(doc, flow_rows)

    # --- Save ---
    doc.save(out_path)
    print(f"[OK] {out_path}  ({Path(out_path).stat().st_size / 1024:.1f} KB)")


# =============================================================================
#  PART 2 -- how_to_run_mbX_Pro.docx (lay-user manual)
# =============================================================================

def build_user_manual(out_path):
    doc = Document()
    set_default_font(doc, name="Calibri", size=11)

    add_logo_centered(doc, width_in=2.6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("How to run mbX Pro")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "A step-by-step guide for non-experts.\n"
        "From a fresh laptop to a publication-ready microbiome report.")
    r.italic = True
    add_page_break(doc)

    # Audience + outcome
    add_h1(doc, "Who this guide is for")
    add_para(doc, (
        "You are someone who has paired-end 16S rRNA FASTQ files plus a "
        "metadata spreadsheet and you want a complete microbiome analysis "
        "report -- including diversity statistics, differential abundance, "
        "machine-learning biomarkers, networks, and figures -- without "
        "running 18 separate command-line tools or writing any code."
    ))
    add_para(doc, "By the end of this guide you will:")
    add_bullet(doc, "Have all dependencies installed correctly on macOS or Linux.")
    add_bullet(doc, "Have the mbX Pro pipeline installed in your $PATH.")
    add_bullet(doc, "Have run a complete analysis with one command.")
    add_bullet(doc, "Have an HTML + PDF report ready to send to a collaborator.")

    # Part 1 -- requirements
    add_h1(doc, "Part 1. Requirements (read this first)")
    add_table_2col(doc, [
        ("Operating system", "macOS 12+ or any modern Linux"),
        ("Disk space",        "10-20 GB free in the directory next to your FASTQ folder"),
        ("RAM",               "8 GB minimum, 16+ GB recommended"),
        ("Internet",          "needed for the first run only (downloads ~3.5 GB of reference databases)"),
        ("Background skill",  "comfort opening a terminal and copy-pasting commands"),
    ])

    # Part 2 -- install dependencies
    add_h1(doc, "Part 2. Install the prerequisites")
    add_h2(doc, "Step 1. Install Miniconda (one-time)")
    add_para(doc, (
        "Miniconda is a small package manager. We will use it to install "
        "QIIME2 in its own isolated environment. If you already have conda "
        "(via Miniconda or Anaconda), skip this step."
    ))
    add_para(doc, "Open Terminal and paste:")
    add_code_block(doc,
        "# macOS (Apple Silicon)\n"
        "curl -O https://repo.anaconda.com/miniconda/Miniconda3-latest-MacOSX-arm64.sh\n"
        "bash Miniconda3-latest-MacOSX-arm64.sh\n"
        "\n"
        "# macOS (Intel)\n"
        "curl -O https://repo.anaconda.com/miniconda/Miniconda3-latest-MacOSX-x86_64.sh\n"
        "bash Miniconda3-latest-MacOSX-x86_64.sh\n"
        "\n"
        "# Linux (x86_64)\n"
        "curl -O https://repo.anaconda.com/miniconda/Miniconda3-latest-Linux-x86_64.sh\n"
        "bash Miniconda3-latest-Linux-x86_64.sh")
    add_para(doc, (
        "Press ENTER and answer 'yes' when prompted. Close and re-open the "
        "Terminal when the installer finishes."
    ))

    add_h2(doc, "Step 2. Install QIIME2 (one-time)")
    add_para(doc, "Run these commands one at a time:")
    add_code_block(doc,
        "# macOS\n"
        "curl -OL https://data.qiime2.org/distro/amplicon/qiime2-amplicon-2025.4-py310-osx-conda.yml\n"
        "conda env create -n qiime2-amplicon-2025.4 --file qiime2-amplicon-2025.4-py310-osx-conda.yml\n"
        "\n"
        "# Linux  (replace 'osx' with 'linux' in the filename)\n"
        "curl -OL https://data.qiime2.org/distro/amplicon/qiime2-amplicon-2025.4-py310-linux-conda.yml\n"
        "conda env create -n qiime2-amplicon-2025.4 --file qiime2-amplicon-2025.4-py310-linux-conda.yml")
    add_para(doc, "This download takes 15-30 minutes depending on your bandwidth.")

    add_h2(doc, "Step 3. Install R (one-time)")
    add_para(doc, "IMPORTANT: install R OUTSIDE conda -- this prevents library conflicts.")
    add_code_block(doc,
        "# macOS  (uses Homebrew; install Homebrew first from https://brew.sh)\n"
        "brew install r\n"
        "\n"
        "# Debian / Ubuntu\n"
        "sudo apt-get install r-base\n"
        "\n"
        "# Fedora / RHEL\n"
        "sudo dnf install R")

    # Part 3 -- install mbX Pro
    add_h1(doc, "Part 3. Install mbX Pro")
    add_h2(doc, "Step 1. Download the mbX Pro repository")
    add_para(doc, "If you have git installed:")
    add_code_block(doc,
        "git clone https://github.com/<author>/mbXPro.git\n"
        "cd mbXPro")
    add_para(doc, "If you don't have git, click the green 'Code' button on the GitHub page, choose 'Download ZIP', then unzip it and `cd` into the unzipped folder.")

    add_h2(doc, "Step 2. Run the installer")
    add_code_block(doc, "bash install/install_mbXPro.sh")
    add_para(doc, "The installer:")
    add_bullet(doc, "Verifies your operating system and required tools.")
    add_bullet(doc, "Copies all 21 step scripts and the logo into ~/bin/.")
    add_bullet(doc, "Adds ~/bin to your PATH (via ~/.zshrc and/or ~/.bashrc).")
    add_bullet(doc, "Prints a 'ready to use' message.")

    add_para(doc, "After the installer finishes, OPEN A NEW TERMINAL (or run `source ~/.zshrc`) and verify:")
    add_code_block(doc, "mbXPro --help")
    add_para(doc, "You should see the help message. If you get 'command not found', re-open Terminal.")

    # Part 4 -- prepare data
    add_h1(doc, "Part 4. Prepare your data")
    add_h2(doc, "Your FASTQ folder")
    add_para(doc, "Put all your raw paired-end FASTQ files (.fastq.gz) into a single folder. The filenames should look like one of these patterns:")
    add_code_block(doc,
        "S1_R1_001.fastq.gz   S1_R2_001.fastq.gz\n"
        "Sample-1_R1.fastq.gz Sample-1_R2.fastq.gz\n"
        "P1_L001_R1_001.fastq.gz   P1_L001_R2_001.fastq.gz")

    add_h2(doc, "Your metadata file")
    add_para(doc, "Create a tab-separated metadata.txt file. Example:")
    add_code_block(doc,
        "sample-id\tsample_type\tsample_density\tsource\n"
        "S1\trumen\thigh\tdairy\n"
        "S2\trumen\tlow\tdairy\n"
        "S3\tfeces\thigh\tbeef\n"
        "...")
    add_para(doc, "Rules:")
    add_bullet(doc, "First column header MUST be 'sample-id' (or 'id', 'sampleid', 'feature-id').")
    add_bullet(doc, "Sample IDs MUST match what mbX Pro extracts from your FASTQ filenames.")
    add_bullet(doc, "No duplicate sample IDs.")
    add_bullet(doc, "Add as many categorical columns as you want (treatment, tissue, batch, ...) -- mbX Pro will run statistics for every one of them.")
    add_para(doc, "(See examples/example_metadata.txt for a working template.)")

    # Part 5 -- run
    add_h1(doc, "Part 5. Run the pipeline (one command)")
    add_code_block(doc,
        "# 1. Activate QIIME2\n"
        "conda activate qiime2-amplicon-2025.4\n"
        "\n"
        "# 2. Run mbX Pro\n"
        "mbXPro /path/to/your/FASTQ /path/to/your/metadata.txt")
    add_para(doc, "That's it. Sit back -- a typical 20-sample run takes 30 to 60 minutes on a laptop.")

    add_h2(doc, "What if my sequencing facility already trimmed primers?")
    add_para(doc,
        "Nothing -- the pipeline handles this automatically. When step 0 cannot "
        "find primers in the reads, it sets DETECTION_STATUS=TRIMMED, identifies "
        "the V-region from conserved 16S anchor motifs, and:")
    add_bullet(doc,
        "Step 3 (DADA2 parameter finder) sets trim-left-f=0 and "
        "trim-left-r=0 -- because the primers are already gone, trimming "
        "another 20 bp would destroy real biological sequence.")
    add_bullet(doc,
        "Step 5 (classifier preparation) automatically switches to "
        "FULL-LENGTH classifier mode: extract-reads is skipped and "
        "Naive-Bayes is trained directly on the entire Greengenes2 "
        "backbone (or downloaded pre-trained from Zenodo, see below).")
    add_bullet(doc,
        "The pipeline finishes normally; the final report clearly states "
        "which mode and which classifier source was used.")
    add_para(doc,
        "If you happen to know the original primer sequences and want to FORCE "
        "region-specific mode (slightly more precise at species level), add:")
    add_code_block(doc,
        "mbXPro /path/to/FASTQ /path/to/metadata.txt \\\n"
        "  --forward-primer CCTACGGGNGGCWGCAG \\\n"
        "  --reverse-primer GACTACHVGGGTATCTAATCC")
    add_para(doc,
        "But you do NOT need to. The default behaviour is self-healing.")

    add_h2(doc, "Where does the full-length classifier come from?")
    add_para(doc,
        "When the pipeline runs in FULL-LENGTH classifier mode (because "
        "primers were already trimmed or could not be detected), step 5 "
        "first tries to DOWNLOAD a pre-trained, sha256-verified Naive-Bayes "
        "classifier from this Zenodo record:")
    add_code_block(doc, "https://zenodo.org/records/20021035")
    add_para(doc,
        "Eight pre-trained classifiers are available, one per supported "
        "QIIME2 release. mbX Pro automatically picks the one matching your "
        "active QIIME2 + scikit-learn version, downloads it (resumable, "
        "sha256-verified), and skips the local training step entirely -- "
        "saving 30-90 minutes per run.")
    add_para(doc,
        "If anything goes wrong (no compatible release for your QIIME2, "
        "network down, sha256 mismatch, or classify-sklearn rejects the "
        "downloaded pickle), the pipeline TRANSPARENTLY falls back to the "
        "old behaviour: it downloads the Greengenes2 backbone and trains "
        "the classifier locally. The pipeline NEVER aborts because of a "
        "Zenodo failure -- the run continues and the final report records "
        "exactly which path was taken (CLASSIFIER_SOURCE = zenodo / cached "
        "/ local-training / local-training-fallback).")
    add_para(doc,
        "If you want to skip the Zenodo attempt entirely (e.g. you are "
        "behind an offline corporate proxy and want to go straight to "
        "local training), pass --skip-zenodo to mbXPro.")

    add_h2(doc, "What if the run fails halfway?")
    add_para(doc, "Re-run the same command with --resume. Already-completed steps will be skipped automatically.")
    add_code_block(doc, "mbXPro /path/to/FASTQ /path/to/metadata.txt --resume")

    # Part 6 -- find your report
    add_h1(doc, "Part 6. Find your report")
    add_para(doc, "When the pipeline finishes, the terminal prints exactly where the report lives. The report is also at this predictable path:")
    add_code_block(doc,
        "<parent of FASTQ>/mbX_pro_outputs_<timestamp>/18_final_report/\n"
        "    mbX_pro_final_report.html   (open in any browser)\n"
        "    mbX_pro_final_report.pdf    (publication-ready, A4)")
    add_para(doc, "Both files are SELF-CONTAINED -- you can email or upload them as-is. No additional files needed.")

    # Part 7 -- troubleshooting
    add_h1(doc, "Part 7. Common issues + fixes")
    add_h3(doc, "'command not found: mbXPro'")
    add_para(doc, "Open a NEW terminal window. If still missing: `source ~/.zshrc` (or your shell's rc file).")

    add_h3(doc, "'qiime: command not found'")
    add_para(doc, "You forgot to activate QIIME2. Run: `conda activate qiime2-amplicon-2025.4`")

    add_h3(doc, "'No FASTQ files found in <directory>'")
    add_para(doc, "Verify your FASTQ folder actually contains *.fastq.gz files. Subdirectories up to 3 levels deep are searched automatically.")

    add_h3(doc, "Primer detection failed")
    add_para(doc,
        "Not actually a failure -- the pipeline auto-falls-back to FULL-LENGTH "
        "classifier mode (no primers needed). The final report will note which "
        "mode was used. If you want region-specific mode anyway, supply the "
        "primers manually with --forward-primer and --reverse-primer (see "
        "Part 5).")

    add_h3(doc, "R packages won't install")
    add_para(doc, "Make sure R is installed SYSTEM-WIDE (NOT inside conda). On macOS: `brew install r`. The pipeline will install required R packages automatically on first run.")

    add_h3(doc, "PDF was not generated, only HTML")
    add_para(doc, "Install Google Chrome (best PDF quality) or wkhtmltopdf (`brew install wkhtmltopdf`). The HTML report is always produced regardless.")

    add_h3(doc, "I want to skip the slowest steps")
    add_para(doc, "Use --skip with comma-separated step IDs:")
    add_code_block(doc,
        "# Skip PICRUSt2 (functional inference) and networks\n"
        "mbXPro /path/to/FASTQ /path/to/metadata.txt --skip 15,17")

    # Part 8 -- citing
    add_h1(doc, "Part 8. Citing mbX Pro")
    add_para(doc, "The dedicated mbX Pro paper is in preparation. Until it appears, please cite:")
    add_para(doc, (
        "Lamichhane U., Lourenco J. (2025). mbX: An R Package for Streamlined "
        "Microbiome Analysis. Stats 8(2):44.  doi:10.3390/stats8020044  "
        "https://www.mdpi.com/2571-905X/8/2/44"
    ), italic=True)
    add_para(doc, "The HTML report itself contains a 'how to cite' box with the full citation.")

    # Part 9 -- support
    add_h1(doc, "Part 9. Where to get help")
    add_para(doc, "Open an issue on the mbX Pro GitHub repository. Please include:")
    add_bullet(doc, "Your operating system + version.")
    add_bullet(doc, "The exact command you ran.")
    add_bullet(doc, "The contents of `_pipeline_log/pipeline_master.log` from your output directory.")

    doc.save(out_path)
    print(f"[OK] {out_path}  ({Path(out_path).stat().st_size / 1024:.1f} KB)")


# =============================================================================
#  PART 3 -- how_to_upload_mbX_Pro.docx (maintainer's GitHub guide)
# =============================================================================

def build_upload_guide(out_path):
    doc = Document()
    set_default_font(doc, name="Calibri", size=11)

    add_logo_centered(doc, width_in=2.4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("How to upload mbX Pro to GitHub")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Internal step-by-step for the maintainer.\n"
        "Read this once before publishing v1.0.0; revisit before each release.")
    r.italic = True
    add_page_break(doc)

    add_h1(doc, "Goal")
    add_para(doc, (
        "Publish the mbXPro/ directory tree to a public GitHub repository so "
        "lay users can clone, install, and run the pipeline. After this guide "
        "the repo will be ready for the world."
    ))

    add_h1(doc, "Pre-flight checklist")
    add_para(doc, "Confirm each of the following before pushing:")
    for item in [
        "All 21 step scripts are inside scripts/ and executable (chmod +x).",
        "documentation/ contains the three .docx files plus images/.",
        "assets/mbX_Pro_icon.png is present.",
        "install/install_mbXPro.sh and uninstall_mbXPro.sh are executable.",
        "README.md, LICENSE, VERSION, CITATION.cff, CHANGELOG.md, .gitignore exist at the root.",
        "examples/example_metadata.txt is present.",
        "No mbX_pro_outputs_*/ directories are present (the .gitignore excludes them, but double-check).",
        "No personal data (FASTQ files, real metadata, /Users/<your name>/... paths) anywhere in the tree.",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "Part 1. One-time GitHub setup")
    add_h2(doc, "1.1 Create a GitHub account (if you don't have one)")
    add_para(doc, "Go to https://github.com and sign up. Choose your username carefully; it will appear in the public URL of the repo.")

    add_h2(doc, "1.2 Install git on your machine")
    add_code_block(doc,
        "# macOS (Apple Silicon and Intel)\n"
        "brew install git\n"
        "\n"
        "# Debian / Ubuntu\n"
        "sudo apt-get install git\n"
        "\n"
        "# Fedora / RHEL\n"
        "sudo dnf install git")

    add_h2(doc, "1.3 Configure git with your identity (one-time)")
    add_code_block(doc,
        'git config --global user.name  "Your Full Name"\n'
        'git config --global user.email "your.email@uga.edu"\n'
        'git config --global init.defaultBranch main')

    add_h2(doc, "1.4 Create a Personal Access Token (PAT)")
    add_para(doc, "GitHub no longer accepts plain passwords for git push. Generate a Personal Access Token (classic):")
    add_bullet(doc, "Open https://github.com/settings/tokens")
    add_bullet(doc, "Click 'Generate new token (classic)'.")
    add_bullet(doc, "Name: 'mbXPro publishing token'.")
    add_bullet(doc, "Expiration: 90 days (renew when prompted).")
    add_bullet(doc, "Scopes: tick the entire 'repo' checkbox.")
    add_bullet(doc, "Click 'Generate token' and COPY THE TOKEN (it is shown only once).")
    add_para(doc, "Paste the token in a secure password manager. You will use it as the 'password' the first time git asks.")

    add_h1(doc, "Part 2. Create the GitHub repository (web UI)")
    add_bullet(doc, "Click '+' (top-right of GitHub) -> 'New repository'.")
    add_bullet(doc, "Repository name: mbXPro")
    add_bullet(doc, "Description: Single-command 16S rRNA microbiome pipeline.")
    add_bullet(doc, "Visibility: Public.")
    add_bullet(doc, "Do NOT check 'Add a README file' / 'Add .gitignore' / 'Add license' (we already have those).")
    add_bullet(doc, "Click 'Create repository'.")
    add_para(doc, "Note the URL that appears -- it looks like:")
    add_code_block(doc, "https://github.com/<your-username>/mbXPro.git")

    add_h1(doc, "Part 3. Push the local mbXPro/ directory to GitHub")
    add_para(doc, "Open Terminal and run, replacing <your-username>:")
    add_code_block(doc,
        "# Navigate to the prepared mbXPro/ directory\n"
        "cd /path/to/mbXPro\n"
        "\n"
        "# Initialise as a git repo (only the first time)\n"
        "git init\n"
        "\n"
        "# Stage everything that the .gitignore allows\n"
        "git add .\n"
        "\n"
        "# First commit\n"
        'git commit -m "mbX Pro v1.0.0 -- initial release"\n'
        "\n"
        "# Tell git where the GitHub remote is\n"
        "git branch -M main\n"
        "git remote add origin https://github.com/<your-username>/mbXPro.git\n"
        "\n"
        "# Push (you will be asked for your GitHub username + PAT as password)\n"
        "git push -u origin main")
    add_para(doc, "When git asks for a password, paste your Personal Access Token (NOT your GitHub website password).")

    add_h1(doc, "Part 4. Tag and publish a release")
    add_para(doc, "A 'release' on GitHub gives users a versioned download (with a downloadable ZIP).")
    add_h2(doc, "4.1 Create the version tag locally")
    add_code_block(doc,
        "# Inside the mbXPro/ directory\n"
        'git tag -a v1.0.0 -m "mbX Pro v1.0.0 -- initial release"\n'
        "git push origin v1.0.0")

    add_h2(doc, "4.2 Convert the tag into a release on GitHub")
    add_bullet(doc, "On the repo page, click 'Releases' (right sidebar).")
    add_bullet(doc, "Click 'Draft a new release'.")
    add_bullet(doc, "Choose tag: v1.0.0 (the one you just pushed).")
    add_bullet(doc, "Release title: 'mbX Pro v1.0.0'.")
    add_bullet(doc, "Description: paste the contents of CHANGELOG.md for this version.")
    add_bullet(doc, "Click 'Publish release'.")
    add_para(doc, "GitHub now automatically attaches a downloadable .zip and .tar.gz of the source.")

    add_h1(doc, "Part 5. Verify the public download path")
    add_para(doc, "Open an incognito browser window (so you are not logged in) and visit:")
    add_code_block(doc, "https://github.com/<your-username>/mbXPro")
    add_para(doc, "Confirm:")
    add_bullet(doc, "The README renders with the logo at the top.")
    add_bullet(doc, "The 'Releases' link in the right sidebar shows v1.0.0.")
    add_bullet(doc, "Clicking on documentation/mbXPro_documentation.docx offers a 'View raw' download link.")
    add_para(doc, "Test the install path on a clean machine (or VM):")
    add_code_block(doc,
        "git clone https://github.com/<your-username>/mbXPro.git\n"
        "cd mbXPro\n"
        "bash install/install_mbXPro.sh\n"
        "mbXPro --help")

    add_h1(doc, "Part 6. After the first publish: workflow for future updates")
    add_para(doc, "When you fix a bug or add a feature:")
    add_code_block(doc,
        '# 1. Edit files\n'
        '# 2. Update VERSION (e.g. 1.0.1)\n'
        '# 3. Update CHANGELOG.md\n'
        '\n'
        'git add .\n'
        'git commit -m "describe what changed in one short line"\n'
        'git push\n'
        '\n'
        '# (optional) tag a new release\n'
        'git tag -a v1.0.1 -m "mbX Pro v1.0.1 -- bug fixes"\n'
        'git push origin v1.0.1')

    add_h1(doc, "Part 7. Optional but recommended polish")
    add_h3(doc, "Add a short repo 'About'")
    add_para(doc, "On the repo home page (top-right gear icon next to 'About'):")
    add_bullet(doc, "Description: 'Single-command 16S rRNA microbiome pipeline (FASTQ -> HTML/PDF report).'")
    add_bullet(doc, "Website: leave blank (or link to your lab page).")
    add_bullet(doc, "Topics: microbiome 16s qiime2 dada2 picrust2 ancombc2 random-forest networks bioinformatics")

    add_h3(doc, "Enable the issue tracker")
    add_para(doc, "Settings -> Features -> tick 'Issues'. Users will report bugs here.")

    add_h3(doc, "Add a CITATION button")
    add_para(doc, "GitHub auto-detects CITATION.cff and adds a 'Cite this repository' button on the right sidebar -- you already shipped CITATION.cff, so this works automatically.")

    add_h3(doc, "Add a CI badge to the README (optional, future)")
    add_para(doc, "If you later add GitHub Actions tests, paste the auto-generated badge into README.md so visitors immediately see test status.")

    add_h1(doc, "Part 8. Sharing the pipeline")
    add_para(doc, "Now that the repo is public, point users at:")
    add_bullet(doc, "The repo URL: https://github.com/<your-username>/mbXPro")
    add_bullet(doc, "The user manual: documentation/how_to_run_mbX_Pro.docx (clickable inside GitHub).")
    add_bullet(doc, "The technical reference: documentation/mbXPro_documentation.docx.")

    add_h1(doc, "Troubleshooting")
    add_h3(doc, "git push asks for password forever")
    add_para(doc, "You typed your GitHub website password instead of your Personal Access Token. Generate a new PAT (Part 1.4) and use that.")

    add_h3(doc, "'fatal: remote origin already exists'")
    add_para(doc, "You already added a remote. Run `git remote set-url origin <new url>` instead of `git remote add`.")

    add_h3(doc, "Files >100 MB get rejected")
    add_para(doc, "GitHub blocks individual files >100 MB. The pipeline does not contain any. If you accidentally added a FASTQ file, run:")
    add_code_block(doc,
        "git rm --cached path/to/big-file\n"
        "echo 'path/to/big-file' >> .gitignore\n"
        'git commit -m "untrack large file"\n'
        "git push")

    add_h3(doc, "I made a typo in the latest commit message")
    add_code_block(doc, 'git commit --amend -m "the corrected message"\ngit push --force-with-lease')

    doc.save(out_path)
    print(f"[OK] {out_path}  ({Path(out_path).stat().st_size / 1024:.1f} KB)")


# =============================================================================
#  Main
# =============================================================================

def main():
    out_dir = HERE
    build_technical_doc(out_dir / "mbXPro_documentation.docx")
    build_user_manual(  out_dir / "how_to_run_mbX_Pro.docx")
    build_upload_guide( out_dir / "how_to_upload_mbX_Pro.docx")


if __name__ == "__main__":
    main()
